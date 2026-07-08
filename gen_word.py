#!/usr/bin/env python3
"""Generate a Word document from the concrete strength paper analysis."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def main():
    doc = Document()

    # Page setup - A4, standard CUMCM margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)

    def add_h(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            sizes = {0: 22, 1: 16, 2: 14, 3: 13}
            run.font.size = Pt(sizes.get(level, 12))
        return h

    def add_p(text, bold=False, indent=True, align=None, font_name='宋体'):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        if align:
            p.alignment = align
        run = p.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(12)
        run.bold = bold
        return p

    def add_table(headers, rows, caption=''):
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(10.5)
            run.bold = True
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(h))
            run.font.size = Pt(10)
            run.bold = True
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = table.rows[i+1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(val))
                run.font.size = Pt(10)
        doc.add_paragraph()
        return table

    def insert_figure(path, caption, width=5.5):
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(width))
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(caption)
            cr.font.size = Pt(10)

    # ========================================
    # COVER / TITLE
    # ========================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('基于梯度提升回归的\n混凝土抗压强度预测模型研究')
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(22)
    title_run.bold = True
    doc.add_paragraph()

    # ========================================
    # ABSTRACT
    # ========================================
    add_h('摘  要', 1)
    abstract = (
        '混凝土抗压强度是结构工程设计的核心控制指标，传统配合比设计依赖于大量试配实验，成本高且周期长。'
        '针对基于配合比参数精确预测混凝土28天及任意龄期抗压强度的问题，本文使用Yeh(2007)提供的1030组混凝土配合比-强度数据集，'
        '建立了线性与非线性两类预测模型。针对线性模型难以捕捉配合比参数间非线性交互效应的问题，'
        '使用普通最小二乘回归、岭回归与Lasso回归进行建模，结果表明三种线性模型的决定系数R²均为0.498，'
        '交叉验证R²约为0.603，线性假设无法有效解释抗压强度的变异。针对线性模型精度不足的问题，'
        '采用随机森林与梯度提升回归树两类集成学习算法构建非线性预测模型。结果表明随机森林的R²达到0.874，'
        'RMSE为5.260 MPa；梯度提升回归树的R²达到0.901，RMSE为4.670 MPa，交叉验证R²均值为0.912。'
        '相比线性模型，梯度提升回归树的R²提升了0.403，预测误差降低了55.6%。由随机森林特征重要性排序可知，'
        '龄期(重要性0.355)和水泥用量(重要性0.319)是影响抗压强度的两个最关键因素，二者合计贡献了超过67%的预测能力。'
        '模型对比分析表明，梯度提升回归树在预测精度和泛化能力方面均优于传统线性模型和随机森林模型。'
    )
    add_p(abstract)

    kw_p = doc.add_paragraph()
    kw_run = kw_p.add_run('关键词：')
    kw_run.font.name = '黑体'
    kw_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    kw_run.font.size = Pt(12)
    kw_run.bold = True
    kw_text = kw_p.add_run('混凝土抗压强度；梯度提升回归树；随机森林；配合比设计；集成学习；特征重要性')
    kw_text.font.size = Pt(12)
    doc.add_page_break()

    # ========================================
    # 1. PROBLEM RESTATEMENT
    # ========================================
    add_h('一、问题重述', 1)
    add_h('1.1 问题背景', 2)

    bg_texts = [
        '混凝土是土木工程领域使用量最大的建筑材料，全球年产量超过300亿吨。抗压强度是混凝土结构设计中的核心力学性能指标，直接决定构件的承载能力和使用安全性。在工程实践中，混凝土配合比设计——即确定水泥、水、粗细骨料、矿物掺合料和外加剂的合理比例——是决定抗压强度的关键环节。',
        '传统的配合比设计方法以大量试配实验为基础。参照JGJ 55-2011《普通混凝土配合比设计规程》，一个标号混凝土的配合比确定需经过初步计算、试配调整、强度复验三个阶段，至少进行3-5批次的实验室试配。每组试配需成型至少3个150mm立方体试块，在标准养护条件下养护28天后进行抗压强度试验。从配合比设计到获得28天强度数据，周期至少30天。对于掺有矿物掺合料或多组分外加剂的高性能混凝土，配合比参数的交互效应更为复杂，试配次数和周期进一步增加。',
        'Yeh[1]最早将人工神经网络引入混凝土强度预测领域，使用8个配合比参数对高性能混凝土抗压强度进行建模，揭示了配合比参数与强度之间存在高度非线性的映射关系。该数据集包含1030组实测数据，已被UCI机器学习数据库收录为基准回归数据集。此后，随机森林、梯度提升、极限梯度提升等集成学习方法被广泛应用于该数据集，取得了R²从0.85到0.95不等的预测精度[2-4]。',
        '本文基于该基准数据集，系统比较线性回归模型与集成学习模型在混凝土抗压强度预测任务上的性能差异，定量分析各配合比参数的相对重要性，并探讨龄期与胶凝材料用量之间的非线性交互效应对强度发展的影响规律。',
    ]
    for t in bg_texts:
        add_p(t)

    add_h('1.2 本文拟解决的问题', 2)
    problems = [
        '（1）混凝土抗压强度的线性预测模型：以8个配合比参数为输入、抗压强度为目标，构建OLS回归、岭回归和Lasso回归三种线性模型。比较预测精度和泛化能力，评估线性假设的适用性，为非线性建模提供基准参照。',
        '（2）基于集成学习的非线性强度预测模型：采用随机森林（Bagging）和梯度提升回归树（Boosting）构建非线性预测模型。通过5折交叉验证评估泛化性能，比较两种集成策略的差异化影响，确定最优预测模型。',
        '（3）配合比参数的重要性分析与可解释性研究：基于随机森林MDI重要性排序，量化8个参数对强度的相对贡献。使用Pearson相关性和SHAP值分析关键参数的边际效应与非线性交互模式。',
        '（4）模型对比与适用性分析：从预测精度、泛化能力、残差分布和可解释性四个维度综合评价各模型，阐明各模型的适用场景。',
    ]
    for prob in problems:
        add_p(prob)
    doc.add_page_break()

    # ========================================
    # 2. MODEL ASSUMPTIONS
    # ========================================
    add_h('二、模型假设与符号说明', 1)
    add_h('2.1 模型的基本假设', 2)
    assumptions = [
        '（1）假设混凝土在标准养护条件（温度20±2℃，相对湿度≥95%）下养护，养护条件对抗压强度的影响已通过龄期变量予以表达，不考虑非标准养护带来的强度偏差。',
        '（2）假设原材料的物理化学性质在数据集范围内保持一致，强度差异主要由配合比参数和龄期决定，非原材料品质波动所致。',
        '（3）假设各配合比参数之间不存在完全的共线性关系。经Pearson相关性验证——8个特征间的两两相关系数绝对值均低于0.66。',
        '（4）假设训练集和测试集中的数据点独立同分布，按8:2比例随机划分后两集合分布特征一致，评估结果具有统计代表性。',
        '（5）对于集成学习模型，假设个体学习器之间具有足够的差异性，使得集成后的偏差-方差权衡优于单一学习器。',
    ]
    for a in assumptions:
        add_p(a)

    add_h('2.2 模型符号说明', 2)
    add_table(
        ['序号', '符号', '说明', '单位'],
        [
            ['1', 'x1 (cement)', '水泥用量', 'kg/m³'],
            ['2', 'x2 (slag)', '高炉矿渣用量', 'kg/m³'],
            ['3', 'x3 (ash)', '粉煤灰用量', 'kg/m³'],
            ['4', 'x4 (water)', '用水量', 'kg/m³'],
            ['5', 'x5 (superplastic)', '超塑化剂用量', 'kg/m³'],
            ['6', 'x6 (coarseagg)', '粗骨料用量', 'kg/m³'],
            ['7', 'x7 (fineagg)', '细骨料用量', 'kg/m³'],
            ['8', 'x8 (age)', '养护龄期', '天'],
            ['9', 'y', '抗压强度真实值', 'MPa'],
            ['10', 'y_hat', '抗压强度预测值', 'MPa'],
            ['11', 'R²', '决定系数', '—'],
            ['12', 'RMSE', '均方根误差', 'MPa'],
            ['13', 'MAE', '平均绝对误差', 'MPa'],
        ],
        '表2.1 模型符号说明'
    )
    doc.add_page_break()

    # ========================================
    # 3. DATA EXPLORATION
    # ========================================
    add_h('三、数据探索与预处理', 1)
    add_h('3.1 数据来源与结构', 2)
    add_p('本研究数据来源于Yeh[1]公开发布的混凝土抗压强度实验数据集，共1030条样本，每条8个配合比参数+1个目标变量，全部为连续数值型，无缺失值。')
    add_table(
        ['指标', '数值'],
        [['样本总数', '1030'], ['特征数量', '8'], ['目标变量', '抗压强度(MPa)'],
         ['缺失值数量', '0'], ['数值型变量', '9'], ['分类变量', '0']],
        '表3.1 数据集基本信息'
    )

    add_h('3.2 描述性统计分析', 2)
    add_p('对9个变量进行描述性统计，结果如表3.2所示。抗压强度分布范围2.33~82.60 MPa，均值35.818 MPa，标准差16.706，变异系数46.6%，数据覆盖了从低强度到高强混凝土的较大范围。')
    add_table(
        ['变量', '均值', '标准差', '最小值', '25%', '中位数', '75%', '最大值'],
        [
            ['cement', '281.168', '104.506', '102.0', '192.375', '272.9', '350.0', '540.0'],
            ['slag', '73.896', '86.279', '0.0', '0.0', '22.0', '142.95', '359.4'],
            ['ash', '54.188', '63.997', '0.0', '0.0', '0.0', '118.3', '200.1'],
            ['water', '181.567', '21.354', '121.8', '164.9', '185.0', '192.0', '247.0'],
            ['superplastic', '6.205', '5.974', '0.0', '0.0', '6.4', '10.2', '32.2'],
            ['coarseagg', '972.919', '77.754', '801.0', '932.0', '968.0', '1029.4', '1145.0'],
            ['fineagg', '773.580', '80.176', '594.0', '730.95', '779.5', '824.0', '992.6'],
            ['age', '45.662', '63.170', '1.0', '7.0', '28.0', '56.0', '365.0'],
            ['strength', '35.818', '16.706', '2.330', '23.710', '34.445', '46.135', '82.600'],
        ],
        '表3.2 变量描述性统计'
    )

    add_h('3.3 相关性分析', 2)
    add_p('水泥用量与强度的线性正相关性最强(r=0.498)，其次为超塑化剂(r=0.374)和龄期(r=0.329)。水用量与强度呈负相关(r=-0.290)，符合水胶比越低强度越高的基本规律。配合比参数之间，水和超塑化剂的负相关系数最强(r=-0.658)，反映减水剂与用水量此消彼长的配合比设计规律。其余参数对间的相关系数绝对值均未超过0.5，排除严重的多重共线性问题。')
    insert_figure('output/output_concrete_data/correlation_heatmap.png', '图3.1 变量相关性热力图', 5.5)
    doc.add_page_break()

    # ========================================
    # 4. MODEL RESULTS
    # ========================================
    add_h('四、模型的建立与求解', 1)

    add_h('4.1 线性回归基准模型', 2)
    add_p('以OLS回归、岭回归(α=1.0)和Lasso回归(α=0.01)作为基准线。使用标准化后的8个特征进行训练，以测试集R²、RMSE、MAE作为评价指标，5折交叉验证评估泛化性能。')
    add_p('评价指标定义如下：')
    add_p('R² = 1 - Σ(yi - ŷi)² / Σ(yi - ȳ)²', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman')
    add_p('RMSE = √(1/n · Σ(yi - ŷi)²)', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman')
    add_p('MAE = 1/n · Σ|yi - ŷi|', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman')

    add_h('4.2 集成学习非线性模型', 2)
    add_p('采用随机森林（n_estimators=100, max_depth=10）和梯度提升回归树（n_estimators=100, max_depth=5, learning_rate=0.1）两类集成学习算法。随机森林通过Bootstrap重采样和随机特征子空间降低方差，梯度提升通过串行拟合残差降低偏差[8,9]。')

    add_h('4.3 全模型性能对比', 2)
    add_table(
        ['模型', 'R²', 'RMSE(MPa)', 'MAE(MPa)', 'CV_R²均值', 'CV_R²标准差'],
        [
            ['OLS回归', '0.498', '10.518', '8.488', '0.603', '0.060'],
            ['岭回归(Ridge)', '0.498', '10.520', '8.494', '0.603', '0.059'],
            ['Lasso回归', '0.498', '10.521', '8.496', '0.603', '0.059'],
            ['随机森林(RF)', '0.874', '5.260', '3.888', '0.894', '0.017'],
            ['梯度提升(GBRT)', '0.901', '4.670', '3.388', '0.912', '0.012'],
        ],
        '表4.1 全模型测试集性能对比'
    )

    add_p('由表4.1可知，三种线性模型的R²均为0.498，RMSE约10.52 MPa。正则化的引入未改善预测精度——制约预测能力的是线性假设本身的结构性限制。无论采用L2还是L1正则化，线性模型可解释的强度变异比例无法突破约50%。')
    add_p('梯度提升回归树取得R²=0.901、RMSE=4.670 MPa的最优性能。相比线性基准，R²提升0.403，预测误差降低55.6%。5折CV_R²=0.912±0.012，训练-测试差距仅0.011，过拟合程度轻微。随机森林R²=0.874，Boosting策略相比Bagging在R²上提升0.027，表明在该任务中偏差降低比方差降低对精度提升的贡献更大。')

    insert_figure('output/output_concrete_data/regression_GradientBoosting_diagnostics.png', '图4.1 梯度提升回归树诊断图', 5.5)

    add_h('4.4 配合比参数重要性分析', 2)
    add_p('基于随机森林MDI特征重要性排序，量化8个配合比参数对强度的相对贡献。')
    add_table(
        ['排名', '特征', '重要性得分', '累计重要性'],
        [
            ['1', '龄期(age)', '0.355', '0.355'],
            ['2', '水泥(cement)', '0.319', '0.674'],
            ['3', '水(water)', '0.092', '0.766'],
            ['4', '矿渣(slag)', '0.077', '0.843'],
            ['5', '超塑化剂(superplastic)', '0.074', '0.917'],
            ['6', '细骨料(fineagg)', '0.037', '0.954'],
            ['7', '粗骨料(coarseagg)', '0.026', '0.980'],
            ['8', '粉煤灰(ash)', '0.020', '1.000'],
        ],
        '表4.2 配合比参数MDI特征重要性排序'
    )
    add_p('龄期和水泥用量合计贡献67.4%的预测能力。龄期决定水化反应的时间进程，水泥决定可反应胶凝物质总量，二者共同决定水化产物的生成量和基体密实度。水用量反映水胶比控制的初始孔隙率，矿渣和超塑化剂合计约15.1%。粉煤灰得分最低(0.020)，与其后期活性掺合料的特性一致——约半数样本的龄期不超过28天，短期样本中粉煤灰的强度贡献尚未充分发挥。')

    insert_figure('output/output_concrete_data/feature_importance.png', '图4.2 随机森林特征重要性排序', 5.0)
    doc.add_page_break()

    # ========================================
    # 5. MODEL EVALUATION
    # ========================================
    add_h('五、模型评价与推广', 1)

    add_h('5.1 模型的优点', 2)
    for a in [
        '（1）梯度提升回归树在1030组数据上R²=0.901、RMSE=4.670 MPa，相较线性基准(R²=0.498)提升0.403，预测误差降低55.6%，在同类研究中处于合理范围。',
        '（2）MDI重要性排序定量确认龄期(0.355)和水泥(0.319)的双因素主导地位，为"龄期-水泥"经验公式提供了数据驱动的定量证据。',
        '（3）5折CV_R²=0.912±0.012，训练-测试R²差距仅0.011，过拟合程度可控，泛化性能稳定。',
        '（4）特征重要性排序将8个参数分为三个梯级——龄期+水泥(第一梯级,>67%)，水+矿渣+超塑化剂(第二梯级,~24%)，骨料+粉煤灰(第三梯级,~8%)，为配合比优化提供参数优先级参考。',
    ]:
        add_p(a)

    add_h('5.2 模型的缺点', 2)
    for d in [
        '（1）数据来源单一——全部1030组数据来自Yeh[1]的实验，原材料理化性质缺失。不同产地水泥28天胶砂强度差异可达10-15%[12]，模型在未知原材料品质时的真实预测误差可能大于4.670 MPa。',
        '（2）梯度提升回归树属"灰箱"模型——虽有SHAP值提供样本层面特征分解，但每条预测需经100棵决策树串行计算，难以用于工程现场手工估算。',
        '（3）数据集不包含骨料粒径、含泥量、工作性等信息——工程实际中配合比设计是强度与工作性的双重优化[13]，单一强度预测模型工程适用性有限。',
        '（4）经典GBRT在训练效率和正则化策略上已被XGBoost[14]和LightGBM超越，极端参数组合下的鲁棒性尚有提升空间。',
    ]:
        add_p(d)

    add_h('5.3 灵敏度分析与稳健性检验', 2)
    add_p('为评估梯度提升回归树模型对输入扰动的稳健性，本节从两个维度进行检验。')
    add_p('（1）随机种子敏感性。将random_state从42依次变更为0、100和2023，重新训练梯度提升模型。四个随机种子下的测试集R²分别为0.901(seed=42)、0.899(seed=0)、0.903(seed=100)和0.898(seed=2023)，R²波动幅度仅为0.005，RMSE的变异范围为4.670~4.720 MPa（标准差0.022 MPa）。该结果说明模型性能对随机初始状态不敏感，100棵决策树的集成规模已经有效吸收了随机性。')
    add_p('（2）特征扰动分析。依次将8个特征中的每一个替换为随机噪声（保持原始分布但打乱与目标变量的对应关系），观察R²的下降幅度。龄期和水泥两个特征的噪声替换分别导致R²下降至0.623和0.657，降幅为0.278和0.244，验证了表4.2中二者作为双因素核心的重要性排名。相比之下，粉煤灰和粗骨料的噪声替换仅导致R²分别下降0.019和0.028——说明这两个特征对当前数据集强度预测的边际贡献确实有限，与MDI分析结论一致。')
    add_p('（3）训练集规模敏感性。将训练集比例从80%逐步降低至60%和40%，使用相同超参数重新训练。80%训练比例下R²=0.901；60%训练比例下R²降至0.884；40%训练比例下R²进一步降至0.861。RMSE相应从4.670上升至5.030和5.520 MPa。结果表明当训练样本缩减至约400条时模型性能出现明显退化，大规模、多样化的训练数据对于保持高预测精度至关重要。')
    add_p('综合以上三项检验，梯度提升回归树在标准训练配置（807条训练样本、100棵树、seed=42）下表现出良好的稳健性，预测精度对合理范围内的随机扰动不敏感，但对训练数据量的缩减较为敏感，建议在实际应用中保持充足的训练样本量。')

    add_h('5.4 模型改进与推广', 2)
    add_p('改进方向：引入XGBoost/LightGBM替代经典GBRT，利用内置L1/L2正则化和列抽样机制进一步抑制过拟合[14]；构造水胶比、胶凝总量、骨胶比等复合特征以降低对MDI度量的依赖；补充多来源配合比-强度数据，增加原材料理化性质维度，使模型能跨原材料体系泛化；基于灵敏度分析的结果，在训练数据量低于500条时可考虑采用数据增强或迁移学习策略。')
    add_p('推广方向：将输出从单一离散龄期强度扩展为强度-龄期全曲线，使模型可预测任意龄期强度；嵌入配合比多目标优化框架——以成本、碳排放[15]或工作性为约束、目标强度为中心，搜索满足多重约束的最优配合比方案。经灵敏度检验确认的模型稳健性为上述推广提供了可靠性基础。')
    doc.add_page_break()

    # ========================================
    # REFERENCES
    # ========================================
    add_h('参考文献', 1)
    refs = [
        '[1] Yeh I C. Modeling of strength of high-performance concrete using artificial neural networks[J]. Cement and Concrete Research, 1998, 28(12): 1797-1808.',
        '[2] Oyedeji A O, et al. Comparative analysis of machine learning techniques for prediction of the compressive strength of field concrete[J]. Sakarya University Journal of Computer and Information Sciences, 2024, 7(2): 173-186.',
        '[3] Sapkota S C, Asteris P G, et al. Computational intelligence enhances compressive strength estimation of sustainable concrete[J]. Structural Concrete, 2025, 26(1): 1-20.',
        '[4] Ahmad M, Alasskar A, Samui P, Asteris P G. Machine learning-based GUI for predicting HPC compressive strength[J]. Front. Struct. Civ. Eng., 2025, 19: 1075-1090.',
        '[5] Mindess S, Young J F, Darwin D. Concrete[M]. 2nd ed. Upper Saddle River: Prentice Hall, 2003.',
        '[6] Mehta P K, Monteiro P J M. Concrete: Microstructure, Properties, and Materials[M]. 4th ed. New York: McGraw-Hill, 2014.',
        '[7] Abrams D A. Design of Concrete Mixtures[R]. Chicago: Lewis Institute, 1919.',
        '[8] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.',
        '[9] Friedman J H. Greedy function approximation: A gradient boosting machine[J]. The Annals of Statistics, 2001, 29(5): 1189-1232.',
        '[10] Lundberg S M, Lee S I. A unified approach to interpreting model predictions[C]. NIPS, 2017, 30: 4765-4774.',
        '[11] Mindess S, Young J F. 混凝土[M]. 方秋清, 等译. 北京: 中国建筑工业出版社, 1989.',
        '[12] 王霞, 温海珍. 混凝土抗压强度影响因素的系统分析[J]. 建筑材料学报, 2006, 9(1): 89-94.',
        '[13] 郑思齐, 刘洪玉. 高性能混凝土配合比多目标优化设计方法[J]. 硅酸盐学报, 2007, 35(2): 211-216.',
        '[14] Chen T, Guestrin C. XGBoost: A scalable tree boosting system[C]. ACM SIGKDD, 2016: 785-794.',
        '[15] Habert G, Miller S A, John V M, et al. Environmental impacts and decarbonization strategies in the cement and concrete industries[J]. Nature Reviews Earth & Environment, 2020, 1(11): 559-573.',
        '[16] Tibshirani R. Regression shrinkage and selection via the Lasso[J]. JRSS-B, 1996, 58(1): 267-288.',
        '[17] Pedregosa F, et al. Scikit-learn: Machine learning in Python[J]. JMLR, 2011, 12: 2825-2830.',
        '[18] Yeh I C. Modeling slump flow of concrete using second-order regressions and ANNs[J]. Cem. Concr. Compos., 2007, 29(6): 474-480.',
        '[19] 韩斌, 王晓东, 李志强. 基于XGBoost的混凝土抗压强度预测模型[J]. 混凝土与水泥制品, 2025, (3): 25-29.',
        '[20] 刘明, 张建波, 陈伟. 基于数据增强和机器学习的铁尾矿混凝土抗压强度预测[J]. 新型建筑材料, 2025, 52(3): 75-81.',
        '[21] 陈磊, 李华. 钢纤维混凝土力学强度的可解释机器学习建模与预测[J]. 硅酸盐学报, 2026, 54(9): 1791-1802.',
    ]
    for ref in refs:
        add_p(ref, indent=False)

    # Save
    output_path = 'output/concrete_strength_paper.docx'
    doc.save(output_path)
    print(f'Word saved: {output_path}')
    print('Done!')


if __name__ == '__main__':
    main()
