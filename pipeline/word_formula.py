#!/usr/bin/env python3
"""
Word 公式模块 — 使用 OMML (Office Math Markup Language) 插入真公式对象
禁止以纯文本 Unicode 写入公式。
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree
import copy

# OMML namespace
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _make_math_run(text, italic=True):
    """Create a math run element <m:r>"""
    r = OxmlElement('m:r')
    if italic:
        rPr = OxmlElement('m:rPr')
        nor = OxmlElement('m:nor')
        rPr.append(nor)
        r.append(rPr)
    t = OxmlElement('m:t')
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def _make_math_text(text):
    """Create a regular text run in math (not italic)"""
    return _make_math_run(text, italic=False)


def _make_math_italic(text):
    """Create an italic math run"""
    r = OxmlElement('m:r')
    t = OxmlElement('m:t')
    t.text = str(text)
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def _make_subsuper(base, sub=None, sup=None):
    """Create subscript/superscript element <m:sSubSup> or <m:sSub>"""
    if sub and sup:
        elem = OxmlElement('m:sSubSup')
        elem.append(_make_math_italic(base))
        e_sub = OxmlElement('m:sub')
        e_sub.append(_make_math_italic(sub))
        elem.append(e_sub)
        e_sup = OxmlElement('m:sup')
        e_sup.append(_make_math_italic(sup))
        elem.append(e_sup)
    elif sub:
        elem = OxmlElement('m:sSub')
        elem.append(_make_math_italic(base))
        e_sub = OxmlElement('m:sub')
        e_sub.append(_make_math_italic(sub))
        elem.append(e_sub)
    elif sup:
        elem = OxmlElement('m:sSup')
        elem.append(_make_math_italic(base))
        e_sup = OxmlElement('m:sup')
        e_sup.append(_make_math_italic(sup))
        elem.append(e_sup)
    else:
        return _make_math_italic(base)
    return elem


def _make_frac(num, den):
    """Create fraction <m:f>"""
    f = OxmlElement('m:f')
    num_e = OxmlElement('m:num')
    num_e.append(_make_math_italic(num) if isinstance(num, str) else num)
    f.append(num_e)
    den_e = OxmlElement('m:den')
    den_e.append(_make_math_italic(den) if isinstance(den, str) else den)
    f.append(den_e)
    return f


def _make_sqrt(inner):
    """Create square root <m:rad>"""
    rad = OxmlElement('m:rad')
    radPr = OxmlElement('m:radPr')
    deg = OxmlElement('m:degHide')
    deg.set(qn('m:val'), '1')
    radPr.append(deg)
    rad.append(radPr)
    e = OxmlElement('m:e')
    e.append(_make_math_italic(inner) if isinstance(inner, str) else inner)
    rad.append(e)
    return rad


def _make_sum():
    """Create summation symbol"""
    return _make_math_italic('∑')  # ∑


def _make_hat(base):
    """Create hat accent <m:acc>"""
    acc = OxmlElement('m:acc')
    accPr = OxmlElement('m:accPr')
    chr_e = OxmlElement('m:chr')
    chr_e.set(qn('m:val'), '̂')
    accPr.append(chr_e)
    acc.append(accPr)
    e = OxmlElement('m:e')
    e.append(_make_math_italic(base) if isinstance(base, str) else base)
    acc.append(e)
    return acc


def _make_bar(base):
    """Create bar accent (for mean)"""
    acc = OxmlElement('m:acc')
    accPr = OxmlElement('m:accPr')
    chr_e = OxmlElement('m:chr')
    chr_e.set(qn('m:val'), '̄')
    accPr.append(chr_e)
    acc.append(accPr)
    e = OxmlElement('m:e')
    e.append(_make_math_italic(base) if isinstance(base, str) else base)
    acc.append(e)
    return acc


def _make_group(*elements):
    """Group multiple elements in a row"""
    if len(elements) == 1:
        return elements[0]
    # For multiple elements, wrap in a d (delimiter) or just return as list
    return list(elements)


def _build_omath(*elements):
    """Build a complete <m:oMath> element"""
    omath = OxmlElement('m:oMath')
    for elem in elements:
        if isinstance(elem, list):
            for e in elem:
                omath.append(e)
        else:
            omath.append(elem)
    return omath


def insert_formula(paragraph, formula_type, number=""):
    """
    Insert a Word equation object into the paragraph.

    Args:
        paragraph: docx Paragraph object
        formula_type: str, one of the preset formula types
        number: str, equation number like "1"

    Preset formula types:
        - "ols": y = β₀ + β₁x₁ + β₂x₂ + ... + ε
        - "r2": R² = 1 - Σ(yᵢ-ŷᵢ)² / Σ(yᵢ-ȳ)²
        - "rmse": RMSE = √(1/n · Σ(yᵢ-ŷᵢ)²)
        - "mae": MAE = 1/n · Σ|yᵢ-ŷᵢ|
        - "zscore": x_std = (x-μ)/σ
        - "argmin": β̂ = argmin Σ(...)
    """
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if formula_type == "ols":
        # y = β₀ + β₁x₁ + β₂x₂ + β₃x₃ + β₄x₄ + ε
        omath = _build_omath(
            _make_math_italic('y'),
            _make_math_run(' = '),
            _make_subsuper('β', sub='0'),
            _make_math_run(' + '),
            _make_subsuper('β', sub='1'),
            _make_math_italic('x'),
            _make_subsuper('', sub='1'),
            _make_math_run(' + '),
            _make_subsuper('β', sub='2'),
            _make_math_italic('x'),
            _make_subsuper('', sub='2'),
            _make_math_run(' + '),
            _make_subsuper('β', sub='3'),
            _make_math_italic('x'),
            _make_subsuper('', sub='3'),
            _make_math_run(' + '),
            _make_subsuper('β', sub='4'),
            _make_math_italic('x'),
            _make_subsuper('', sub='4'),
            _make_math_run(' + '),
            _make_math_italic('ε'),
        )
    elif formula_type == "r2":
        # R² = 1 - Σ(yᵢ-ŷᵢ)²/Σ(yᵢ-ȳ)²
        omath = _build_omath(
            _make_subsuper('R', sup='2'),
            _make_math_run(' = 1 - '),
            _make_frac(
                _build_omath(
                    _make_sum(),
                    _make_math_run('('),
                    _make_math_italic('y'),
                    _make_subsuper('', sub='i'),
                    _make_math_run(' - '),
                    _make_hat('y'),
                    _make_subsuper('', sub='i'),
                    _make_math_run(')'),
                    _make_subsuper('', sup='2'),
                ),
                _build_omath(
                    _make_sum(),
                    _make_math_run('('),
                    _make_math_italic('y'),
                    _make_subsuper('', sub='i'),
                    _make_math_run(' - '),
                    _make_bar('y'),
                    _make_math_run(')'),
                    _make_subsuper('', sup='2'),
                ),
            )
        )
    elif formula_type == "rmse":
        # RMSE = √(1/n · Σ(yᵢ-ŷᵢ)²)
        omath = _build_omath(
            _make_math_text('RMSE'),
            _make_math_run(' = '),
            _make_sqrt(
                _build_omath(
                    _make_frac('1', 'n'),
                    _make_math_run(' · '),
                    _make_sum(),
                    _make_math_run('('),
                    _make_math_italic('y'),
                    _make_subsuper('', sub='i'),
                    _make_math_run(' - '),
                    _make_hat('y'),
                    _make_subsuper('', sub='i'),
                    _make_math_run(')'),
                    _make_subsuper('', sup='2'),
                )
            )
        )
    elif formula_type == "mae":
        # MAE = 1/n · Σ|yᵢ-ŷᵢ|
        omath = _build_omath(
            _make_math_text('MAE'),
            _make_math_run(' = '),
            _make_frac('1', 'n'),
            _make_math_run(' · '),
            _make_sum(),
            _make_math_run('|'),
            _make_math_italic('y'),
            _make_subsuper('', sub='i'),
            _make_math_run(' - '),
            _make_hat('y'),
            _make_subsuper('', sub='i'),
            _make_math_run('|'),
        )
    elif formula_type == "zscore":
        # x_std = (x-μ)/σ
        omath = _build_omath(
            _make_subsuper('x', sub='std'),
            _make_math_run(' = '),
            _make_frac(
                _build_omath(
                    _make_math_italic('x'),
                    _make_math_run(' - '),
                    _make_math_italic('μ'),
                ),
                _make_math_italic('σ'),
            )
        )
    elif formula_type == "argmin":
        # β̂ = argmin Σ(yᵢ - xᵢᵀβ)²
        omath = _build_omath(
            _make_hat('β'),
            _make_math_run(' = argmin '),
            _make_sum(),
            _make_math_run('('),
            _make_math_italic('y'),
            _make_subsuper('', sub='i'),
            _make_math_run(' - '),
            _make_math_italic('x'),
            _make_subsuper('', sub='i'),
            _make_subsuper('', sup='T'),
            _make_math_italic('β'),
            _make_math_run(')'),
            _make_subsuper('', sup='2'),
        )
    else:
        # Generic: treat as plain LaTeX-style and write as text for now
        omath = _build_omath(_make_math_italic(formula_type))

    paragraph._element.append(omath)

    if number:
        run = paragraph.add_run(f"    ({number})")
        run.font.size = Pt(12)


# ============================================================
# 便利函数：在文档中插入公式段落
# ============================================================
def add_formula_paragraph(doc, formula_type, number=""):
    """在文档末尾添加一个带公式的段落"""
    p = doc.add_paragraph()
    insert_formula(p, formula_type, number)
    return p


# ============================================================
# 兼容旧版 FMLA() 函数签名
# ============================================================
def fmla_omml(doc, formula_type, number=""):
    """兼容 gen_v4.py 的 FMLA() 调用"""
    return add_formula_paragraph(doc, formula_type, number)


if __name__ == "__main__":
    # 测试：生成一个包含真Word公式的文档
    doc = Document()
    doc.add_heading("Word OMML Formula Test", level=1)

    doc.add_paragraph("OLS Regression Equation:")
    add_formula_paragraph(doc, "ols", "1")

    doc.add_paragraph("R-squared:")
    add_formula_paragraph(doc, "r2", "2")

    doc.add_paragraph("Root Mean Square Error:")
    add_formula_paragraph(doc, "rmse", "3")

    doc.add_paragraph("Mean Absolute Error:")
    add_formula_paragraph(doc, "mae", "4")

    doc.add_paragraph("Z-score Standardization:")
    add_formula_paragraph(doc, "zscore", "5")

    doc.add_paragraph("OLS Estimator:")
    add_formula_paragraph(doc, "argmin", "6")

    out = "C:/Users/HZC12/Desktop/test_formula.docx"
    doc.save(out)
    print(f"Saved: {out}")
    print("Open in Word and double-click any formula to verify it's a real equation object.")
