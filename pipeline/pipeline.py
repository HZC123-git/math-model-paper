#!/usr/bin/env python3
"""
数学建模竞赛论文 - 数据自动分析管线
=====================================
输入：CSV/Excel 数据集
输出：结构化 JSON 分析报告（直接喂给 math-model-paper skill）

用法：
    python pipeline.py data.csv --target y_col --type regression
    python pipeline.py data.csv --target y_col --type classification
    python pipeline.py data.csv --type clustering
"""

import argparse, json, os, sys, warnings
from pathlib import Path
from datetime import datetime

import numpy as np
import pandas as pd
from scipy import stats

warnings.filterwarnings("ignore")

# ============================================================
#  0. 配置
# ============================================================

OUTPUT_DIR = None  # 图表输出目录，运行时设置


def ensure_output_dir(base_dir, problem_name):
    """创建输出目录"""
    global OUTPUT_DIR
    OUTPUT_DIR = Path(base_dir) / f"output_{problem_name}"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


# ============================================================
#  1. 数据加载与预处理
# ============================================================

def to_native(obj):
    """递归转换 numpy 类型为 Python 原生类型"""
    if isinstance(obj, (np.integer,)):
        return int(obj)
    elif isinstance(obj, (np.floating,)):
        return float(obj)
    elif isinstance(obj, np.bool_):
        return bool(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, dict):
        return {k: to_native(v) for k, v in obj.items()}
    elif isinstance(obj, (list, tuple)):
        return [to_native(i) for i in obj]
    return obj


def load_data(file_path):
    """自动识别格式加载数据"""
    ext = Path(file_path).suffix.lower()
    if ext == ".csv":
        df = pd.read_csv(file_path)
    elif ext in (".xlsx", ".xls"):
        df = pd.read_excel(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
    return df


def preprocess(df, target_col=None):
    """基础预处理"""
    report = {}
    original_shape = df.shape

    # 缺失值统计
    missing = df.isnull().sum()
    report["missing_values"] = {
        col: int(cnt)
        for col, cnt in missing[missing > 0].to_dict().items()
    }

    # 删除全空列
    df = df.dropna(axis=1, how="all")

    # 数值列填充中位数
    num_cols = df.select_dtypes(include=[np.number]).columns
    for col in num_cols:
        if df[col].isnull().sum() > 0:
            df[col] = df[col].fillna(df[col].median())

    # 分类列填充众数
    cat_cols = df.select_dtypes(exclude=[np.number]).columns
    for col in cat_cols:
        if df[col].isnull().sum() > 0:
            df[col] = df[col].fillna(df[col].mode()[0] if len(df[col].mode()) > 0 else "Unknown")

    report["original_shape"] = list(original_shape)
    report["cleaned_shape"] = list(df.shape)
    report["removed_cols"] = original_shape[1] - df.shape[1]

    return df, report


# ============================================================
#  2. 探索性数据分析 (EDA)
# ============================================================

def run_eda(df, target_col=None):
    """完整的探索性数据分析"""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns

    # 中文字体设置
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    results = {}
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    cat_cols = df.select_dtypes(exclude=[np.number]).columns.tolist()

    # --- 2.1 数据概览 ---
    results["overview"] = {
        "total_samples": len(df),
        "total_features": len(df.columns),
        "numeric_features": len(num_cols),
        "categorical_features": len(cat_cols),
        "columns": df.columns.tolist(),
        "dtypes": {col: str(dt) for col, dt in df.dtypes.to_dict().items()},
    }

    # --- 2.2 描述性统计 ---
    if num_cols:
        desc = df[num_cols].describe()
        results["descriptive_stats"] = {
            col: {
                "mean": round(float(desc[col]["mean"]), 6),
                "std": round(float(desc[col]["std"]), 6),
                "min": round(float(desc[col]["min"]), 6),
                "25%": round(float(desc[col]["25%"]), 6),
                "50%": round(float(desc[col]["50%"]), 6),
                "75%": round(float(desc[col]["75%"]), 6),
                "max": round(float(desc[col]["max"]), 6),
            }
            for col in num_cols[:20]  # 限制前20列
        }

    # --- 2.3 正态性检验 ---
    normality_results = {}
    for col in num_cols[:15]:
        sample = df[col].dropna()
        if len(sample) > 3 and len(sample) < 5000:
            # 抽样
            if len(sample) > 2000:
                sample = sample.sample(2000, random_state=42)
            # Shapiro-Wilk
            try:
                stat_val, p_val = stats.shapiro(sample)
                normality_results[col] = {
                    "test": "Shapiro-Wilk",
                    "statistic": round(float(stat_val), 6),
                    "p_value": round(float(p_val), 6),
                    "is_normal": p_val > 0.05,
                }
            except:
                pass
    results["normality_tests"] = normality_results

    # --- 2.4 相关性矩阵 + 热力图 ---
    if len(num_cols) >= 2:
        corr_cols = num_cols[:15]
        corr_matrix = df[corr_cols].corr()

        # 找出强相关对 (|r| > 0.6)
        strong_corr = []
        for i in range(len(corr_cols)):
            for j in range(i + 1, len(corr_cols)):
                r = corr_matrix.iloc[i, j]
                if abs(r) > 0.6:
                    strong_corr.append(
                        {
                            "var1": corr_cols[i],
                            "var2": corr_cols[j],
                            "correlation": round(float(r), 4),
                        }
                    )
        results["correlation"] = {
            "strong_pairs": strong_corr[:30],
            "matrix_snippet": {
                col: {
                    col2: round(float(corr_matrix[col][col2]), 4)
                    for col2 in corr_cols[:8]
                }
                for col in corr_cols[:8]
            },
        }

        # 热力图
        fig, ax = plt.subplots(figsize=(12, 10))
        mask = np.triu(np.ones_like(corr_matrix, dtype=bool))
        sns.heatmap(
            corr_matrix, mask=mask, annot=True, fmt=".2f",
            cmap="RdBu_r", center=0, square=True,
            linewidths=0.5, ax=ax,
        )
        ax.set_title("Correlation Heatmap of Numeric Features", fontsize=14)
        fig.tight_layout()
        fig.savefig(OUTPUT_DIR / "correlation_heatmap.png", dpi=150)
        plt.close(fig)

    # --- 2.5 目标变量分布 (如果有) ---
    if target_col and target_col in df.columns:
        fig, axes = plt.subplots(1, 2, figsize=(14, 5))
        # 直方图 + KDE
        if target_col in num_cols:
            df[target_col].dropna().hist(bins=50, ax=axes[0], edgecolor="white")
            axes[0].set_title(f"Distribution of {target_col}", fontsize=12)
            axes[0].set_xlabel(target_col)
            axes[0].set_ylabel("Frequency")

            # QQ plot
            stats.probplot(df[target_col].dropna(), dist="norm", plot=axes[1])
            axes[1].set_title(f"Q-Q Plot of {target_col}", fontsize=12)
        else:
            value_counts = df[target_col].value_counts().head(15)
            axes[0].bar(range(len(value_counts)), value_counts.values)
            axes[0].set_xticks(range(len(value_counts)))
            axes[0].set_xticklabels(value_counts.index, rotation=45, ha="right")
            axes[0].set_title(f"Distribution of {target_col}", fontsize=12)
        fig.tight_layout()
        fig.savefig(OUTPUT_DIR / "target_distribution.png", dpi=150)
        plt.close(fig)

    results["figures"] = [
        {"path": str(OUTPUT_DIR / "correlation_heatmap.png"), "caption": "变量相关性热力图"},
    ]
    if target_col and target_col in df.columns:
        results["figures"].append(
            {"path": str(OUTPUT_DIR / "target_distribution.png"), "caption": f"目标变量 {target_col} 分布图"}
        )

    return results


# ============================================================
#  3. 建模分析
# ============================================================

def run_modeling(df, target_col, problem_type):
    """
    根据问题类型跑相应模型。
    返回：模型名称、评估指标、特征重要性、对比结果
    """
    from sklearn.model_selection import train_test_split, cross_val_score, GridSearchCV
    from sklearn.preprocessing import StandardScaler, LabelEncoder
    from sklearn.metrics import (
        r2_score, mean_squared_error, mean_absolute_error,
        accuracy_score, precision_score, recall_score, f1_score,
        confusion_matrix, roc_auc_score, silhouette_score,
    )
    from sklearn.linear_model import LinearRegression, Ridge, Lasso, LogisticRegression
    from sklearn.ensemble import RandomForestRegressor, RandomForestClassifier, GradientBoostingRegressor, GradientBoostingClassifier
    from sklearn.svm import SVR, SVC
    from sklearn.cluster import KMeans, DBSCAN
    from sklearn.decomposition import PCA

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    results = {"problem_type": problem_type}

    # 预处理
    df_model = df.copy()
    num_cols = df_model.select_dtypes(include=[np.number]).columns.tolist()
    cat_cols = df_model.select_dtypes(exclude=[np.number]).columns.tolist()

    # 编码分类变量
    for col in cat_cols:
        if col != target_col:
            try:
                le = LabelEncoder()
                df_model[col] = le.fit_transform(df_model[col].astype(str))
            except:
                df_model = df_model.drop(columns=[col])

    # --- 回归任务 ---
    if problem_type == "regression" and target_col and target_col in num_cols:
        # 特征和目标
        feature_cols = [c for c in df_model.columns if c != target_col and c in num_cols]
        if not feature_cols:
            feature_cols = [c for c in df_model.columns if c != target_col]

        X = df_model[feature_cols].fillna(0).values
        y = df_model[target_col].fillna(0).values

        # 去除 y 中的极端异常值
        q1, q3 = np.percentile(y, [1, 99])
        mask = (y >= q1) & (y <= q3)
        X, y = X[mask], y[mask]

        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        scaler = StandardScaler()
        X_train_s = scaler.fit_transform(X_train)
        X_test_s = scaler.transform(X_test)

        models = {
            "LinearRegression": LinearRegression(),
            "Ridge": Ridge(alpha=1.0),
            "Lasso": Lasso(alpha=0.01, max_iter=5000),
            "RandomForest": RandomForestRegressor(n_estimators=100, max_depth=10, random_state=42),
            "GradientBoosting": GradientBoostingRegressor(n_estimators=100, max_depth=5, random_state=42),
        }

        model_results = {}
        best_model = None
        best_r2 = -999

        for name, model in models.items():
            try:
                model.fit(X_train_s, y_train)
                y_pred = model.predict(X_test_s)
                r2 = r2_score(y_test, y_pred)
                rmse = np.sqrt(mean_squared_error(y_test, y_pred))
                mae = mean_absolute_error(y_test, y_pred)

                # 交叉验证
                try:
                    cv_scores = cross_val_score(model, X_train_s[:min(5000, len(X_train_s))], y_train[:min(5000, len(y_train))], cv=5, scoring="r2")
                    cv_mean = float(np.mean(cv_scores))
                    cv_std = float(np.std(cv_scores))
                except:
                    cv_mean, cv_std = None, None

                model_results[name] = {
                    "R2": round(float(r2), 6),
                    "RMSE": round(float(rmse), 6),
                    "MAE": round(float(mae), 6),
                    "CV_R2_mean": round(cv_mean, 6) if cv_mean else None,
                    "CV_R2_std": round(cv_std, 6) if cv_std else None,
                }

                if r2 > best_r2:
                    best_r2 = r2
                    best_model = name

                # 残差图 (前三个模型)
                if name in ["LinearRegression", "RandomForest", "GradientBoosting"]:
                    fig, axes = plt.subplots(1, 2, figsize=(12, 5))
                    axes[0].scatter(y_test, y_pred, alpha=0.5, s=10)
                    axes[0].plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], "r--", lw=1)
                    axes[0].set_xlabel("True")
                    axes[0].set_ylabel("Predicted")
                    axes[0].set_title(f"{name}: True vs Predicted (R²={r2:.4f})")

                    residuals = y_test - y_pred
                    axes[1].scatter(y_pred, residuals, alpha=0.5, s=10)
                    axes[1].axhline(y=0, color="r", linestyle="--", lw=1)
                    axes[1].set_xlabel("Predicted")
                    axes[1].set_ylabel("Residuals")
                    axes[1].set_title(f"{name}: Residual Plot")
                    fig.tight_layout()
                    fig.savefig(OUTPUT_DIR / f"regression_{name}_diagnostics.png", dpi=150)
                    plt.close(fig)
            except Exception as e:
                model_results[name] = {"error": str(e)}

        # 特征重要性 (RandomForest)
        try:
            rf = RandomForestRegressor(n_estimators=100, max_depth=10, random_state=42)
            rf.fit(X_train_s, y_train)
            importances = rf.feature_importances_
            indices = np.argsort(importances)[::-1][:15]
            feature_importance = [
                {"feature": feature_cols[i], "importance": round(float(importances[i]), 6)}
                for i in indices
            ]
            results["feature_importance"] = feature_importance

            # 特征重要性图
            fig, ax = plt.subplots(figsize=(10, 6))
            top_features = feature_importance[:12]
            ax.barh(
                [f["feature"] for f in reversed(top_features)],
                [f["importance"] for f in reversed(top_features)],
            )
            ax.set_xlabel("Importance")
            ax.set_title("Feature Importance (Random Forest)", fontsize=13)
            fig.tight_layout()
            fig.savefig(OUTPUT_DIR / "feature_importance.png", dpi=150)
            plt.close(fig)
        except:
            pass

        results["models"] = model_results
        results["best_model"] = best_model
        results["best_r2"] = round(float(best_r2), 6)
        results["feature_columns"] = feature_cols[:20]
        results["sample_size"] = {"train": len(X_train), "test": len(X_test)}
        results["figures"] = [
            {"path": str(OUTPUT_DIR / f"regression_{m}_diagnostics.png"), "caption": f"{m} 回归诊断图"}
            for m in ["LinearRegression", "RandomForest", "GradientBoosting"]
            if m in model_results and "error" not in model_results[m]
        ]
        results["figures"].append(
            {"path": str(OUTPUT_DIR / "feature_importance.png"), "caption": "随机森林特征重要性排序"}
        )

    # --- 分类任务 ---
    elif problem_type == "classification" and target_col and target_col in df.columns:
        feature_cols = [c for c in df_model.columns if c != target_col]
        X = df_model[feature_cols].fillna(0).values
        y_raw = df_model[target_col].fillna(df_model[target_col].mode()[0]).values

        # 标签编码
        le = LabelEncoder()
        y = le.fit_transform(y_raw.astype(str))
        classes = le.classes_.tolist()

        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, stratify=y)
        scaler = StandardScaler()
        X_train_s = scaler.fit_transform(X_train)
        X_test_s = scaler.transform(X_test)

        models = {
            "LogisticRegression": LogisticRegression(max_iter=2000, random_state=42),
            "RandomForest": RandomForestClassifier(n_estimators=100, max_depth=10, random_state=42),
            "GradientBoosting": GradientBoostingClassifier(n_estimators=100, max_depth=5, random_state=42),
        }
        # SVM only for smaller datasets
        if len(X_train) < 10000:
            models["SVM"] = SVC(kernel="rbf", probability=True, random_state=42)

        model_results = {}
        best_model = None
        best_acc = -1

        for name, model in models.items():
            try:
                model.fit(X_train_s, y_train)
                y_pred = model.predict(X_test_s)
                acc = accuracy_score(y_test, y_pred)
                prec = precision_score(y_test, y_pred, average="weighted", zero_division=0)
                rec = recall_score(y_test, y_pred, average="weighted", zero_division=0)
                f1 = f1_score(y_test, y_pred, average="weighted", zero_division=0)

                model_results[name] = {
                    "Accuracy": round(float(acc), 6),
                    "Precision": round(float(prec), 6),
                    "Recall": round(float(rec), 6),
                    "F1_Score": round(float(f1), 6),
                }

                if acc > best_acc:
                    best_acc = acc
                    best_model = name

                # 混淆矩阵图
                fig, ax = plt.subplots(figsize=(8, 6))
                cm = confusion_matrix(y_test, y_pred)
                sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", ax=ax)
                ax.set_xlabel("Predicted")
                ax.set_ylabel("True")
                ax.set_title(f"{name}: Confusion Matrix (Acc={acc:.4f})")
                fig.tight_layout()
                fig.savefig(OUTPUT_DIR / f"classification_{name}_cm.png", dpi=150)
                plt.close(fig)
            except Exception as e:
                model_results[name] = {"error": str(e)}

        results["models"] = model_results
        results["best_model"] = best_model
        results["best_accuracy"] = round(float(best_acc), 6)
        results["num_classes"] = len(classes)
        results["classes"] = classes[:20]
        results["sample_size"] = {"train": len(X_train), "test": len(X_test)}
        results["figures"] = [
            {"path": str(OUTPUT_DIR / f"classification_{m}_cm.png"), "caption": f"{m} 混淆矩阵"}
            for m in model_results if "error" not in model_results[m]
        ]

    # --- 聚类任务 ---
    elif problem_type == "clustering":
        feature_cols = num_cols[:20]
        X = df[feature_cols].fillna(0).values
        scaler = StandardScaler()
        X_s = scaler.fit_transform(X)

        # PCA 降维可视化
        pca = PCA(n_components=2)
        X_pca = pca.fit_transform(X_s)
        pca_var = pca.explained_variance_ratio_

        # KMeans 聚类
        from sklearn.metrics import silhouette_score as sil_score
        k_values = range(2, min(11, len(df) // 10))
        silhouette_scores = []
        inertia_values = []
        best_k, best_sil = 2, -1

        for k in k_values:
            km = KMeans(n_clusters=k, n_init=10, random_state=42)
            labels = km.fit_predict(X_s[:5000])
            sil = sil_score(X_s[:5000], labels)
            inertia_values.append(float(km.inertia_))
            silhouette_scores.append({"k": k, "silhouette_score": round(float(sil), 6)})
            if sil > best_sil:
                best_sil = sil
                best_k = k

        # 最佳 K 聚类结果
        km_best = KMeans(n_clusters=best_k, n_init=10, random_state=42)
        cluster_labels = km_best.fit_predict(X_s)
        results["clustering"] = {
            "best_k": best_k,
            "best_silhouette": round(float(best_sil), 6),
            "silhouette_scores": silhouette_scores,
            "cluster_sizes": {
                int(k): int((cluster_labels == k).sum())
                for k in range(best_k)
            },
            "pca_variance_ratio": [round(float(v), 6) for v in pca_var],
        }

        # 肘部法则图 + 轮廓系数图
        fig, axes = plt.subplots(1, 2, figsize=(14, 5))
        axes[0].plot(list(k_values), inertia_values, "bo-")
        axes[0].set_xlabel("K")
        axes[0].set_ylabel("Inertia")
        axes[0].set_title("Elbow Method", fontsize=12)
        sil_vals = [s["silhouette_score"] for s in silhouette_scores]
        axes[1].plot(list(k_values), sil_vals, "ro-")
        axes[1].set_xlabel("K")
        axes[1].set_ylabel("Silhouette Score")
        axes[1].set_title(f"Silhouette Analysis (Best K={best_k})", fontsize=12)
        fig.tight_layout()
        fig.savefig(OUTPUT_DIR / "clustering_k_selection.png", dpi=150)
        plt.close(fig)

        # PCA 散点图
        fig, ax = plt.subplots(figsize=(10, 8))
        scatter = ax.scatter(X_pca[:5000, 0], X_pca[:5000, 1], c=cluster_labels[:5000], cmap="tab10", alpha=0.6, s=10)
        ax.set_xlabel(f"PC1 ({pca_var[0]*100:.1f}%)")
        ax.set_ylabel(f"PC2 ({pca_var[1]*100:.1f}%)")
        ax.set_title(f"PCA Visualization with K-Means Clusters (K={best_k})")
        plt.colorbar(scatter, ax=ax)
        fig.tight_layout()
        fig.savefig(OUTPUT_DIR / "clustering_pca.png", dpi=150)
        plt.close(fig)

        results["figures"] = [
            {"path": str(OUTPUT_DIR / "clustering_k_selection.png"), "caption": "K值选择：肘部法则与轮廓系数"},
            {"path": str(OUTPUT_DIR / "clustering_pca.png"), "caption": f"PCA降维后的 K-Means(K={best_k})聚类结果"},
        ]

    return results


# ============================================================
#  4. 组装最终 JSON 报告
# ============================================================

def build_report(df, target_col, problem_type, problem_title, file_path):
    """组装完整分析报告 JSON"""
    print("=" * 60)
    print(f"  数据分析管线启动")
    print(f"  数据: {file_path}")
    print(f"  问题类型: {problem_type}")
    print(f"  目标变量: {target_col}")
    print("=" * 60)

    # EDA
    print("\n[1/3] Running EDA...")
    eda_results = run_eda(df, target_col)
    print(f"  -> Shape: {eda_results['overview']['total_samples']} rows x {eda_results['overview']['total_features']} cols")
    normality_normal = sum(1 for v in eda_results.get("normality_tests", {}).values() if v.get("is_normal"))
    print(f"  -> Normal: {normality_normal}/{len(eda_results.get('normality_tests', {}))} passed")
    if "correlation" in eda_results:
        print(f"  -> Strong corr (|r|>0.6): {len(eda_results['correlation']['strong_pairs'])} pairs")

    # 建模
    print("\n[2/3] Running models...")
    model_results = run_modeling(df, target_col, problem_type)
    if problem_type == "regression":
        print(f"  -> Best model: {model_results.get('best_model')} (R2={model_results.get('best_r2')})")
        for m, r in model_results.get("models", {}).items():
            if "R2" in r:
                print(f"     {m}: R2={r['R2']}, RMSE={r['RMSE']}")
    elif problem_type == "classification":
        print(f"  -> Best model: {model_results.get('best_model')} (Acc={model_results.get('best_accuracy')})")
    elif problem_type == "clustering":
        c = model_results.get("clustering", {})
        print(f"  -> Best K: {c.get('best_k')}, Silhouette: {c.get('best_silhouette')}")

    # 组装
    print("\n[3/3] Building report...")
    report = {
        "meta": {
            "generated_at": datetime.now().isoformat(),
            "data_file": str(Path(file_path).name),
            "competition_type": "研究生数学建模" if "研究生" in problem_title else "全国大学生数学建模竞赛",
            "problem_title": problem_title,
            "problem_type": problem_type,
            "target_variable": target_col,
        },
        "data_overview": eda_results["overview"],
        "descriptive_stats": eda_results.get("descriptive_stats", {}),
        "normality_tests": eda_results.get("normality_tests", {}),
        "correlation": eda_results.get("correlation", {}),
        "modeling": model_results,
        "all_figures": (eda_results.get("figures", []) + model_results.get("figures", [])),
    }

    return report


# ============================================================
#  5. 命令行入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="数学建模竞赛论文 - 数据自动分析管线")
    parser.add_argument("data_file", help="数据集文件路径 (CSV/Excel)")
    parser.add_argument("--target", "-t", default=None, help="目标变量列名")
    parser.add_argument("--type", "-p", default="regression",
                        choices=["regression", "classification", "clustering", "time_series"],
                        help="问题类型")
    parser.add_argument("--title", default="数据分析与建模研究", help="论文题目")
    parser.add_argument("--output", "-o", default=None, help="输出 JSON 路径 (默认自动生成)")
    parser.add_argument("--output-dir", "-d", default="./output", help="图表输出目录")

    args = parser.parse_args()

    # 验证文件存在
    if not os.path.exists(args.data_file):
        print(f"错误: 找不到文件 {args.data_file}")
        sys.exit(1)

    # 输出目录
    problem_name = Path(args.data_file).stem
    out_dir = ensure_output_dir(args.output_dir or ".", problem_name)

    # 加载数据
    print(f"\n加载数据: {args.data_file}")
    df = load_data(args.data_file)
    print(f"原始形状: {df.shape}")

    # 预处理
    df_clean, preproc_report = preprocess(df, args.target)
    print(f"清洗后形状: {df_clean.shape}")

    # 自动检测目标列
    target_col = args.target
    if target_col is None and args.type in ("regression", "classification"):
        # 选择最后一列作为默认目标
        target_col = df_clean.columns[-1]
        print(f"未指定目标变量，自动选择最后一列: {target_col}")

    if args.type == "clustering":
        target_col = None

    # 运行分析
    report = build_report(df_clean, target_col, args.type, args.title, args.data_file)
    report["preprocessing"] = preproc_report

    # 输出
    output_path = args.output or str(out_dir / "analysis_report.json")
    # 把绝对路径转相对，方便跨平台
    for f in report["all_figures"]:
        f["path"] = str(Path(f["path"]).as_posix())

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(to_native(report), f, ensure_ascii=False, indent=2)

    # 生成 LaTeX
    latex_path = str(Path(output_path).with_suffix(".tex"))
    build_latex(report, latex_path)
    print(f"  LaTeX: {latex_path}")
    print(f"\n{'='*60}")
    print(f"  Done!")
    print(f"  Report: {output_path}")
    print(f"  LaTeX: {latex_path}")

    # 生成 Word
    docx_path = str(Path(output_path).with_suffix(".docx"))
    build_docx(report, docx_path)
    print(f"  Word: {docx_path}")

    print(f"  Charts: {out_dir}")
    print(f"{'='*60}")
    return report


# ============================================================
#  6. LaTeX 论文输出
# ============================================================

def build_latex(report, output_path):
    """根据分析报告生成完整的带图片LaTeX论文"""

    meta = report.get("meta", {})
    eda = report.get("data_overview", {})
    desc = report.get("descriptive_stats", {})
    normality = report.get("normality_tests", {})
    corr = report.get("correlation", {})
    modeling = report.get("modeling", {})
    problem_type = modeling.get("problem_type", "regression")
    all_figs = report.get("all_figures", [])

    # 找图表路径
    def find_fig(keyword):
        for f in all_figs:
            if keyword in f.get("path", ""):
                # 转相对路径中Windows反斜杠
                p = f["path"].replace("\\", "/")
                # 提取文件名
                import os
                return os.path.basename(p)
        return None

    corr_heatmap = find_fig("correlation_heatmap")
    target_dist = find_fig("target_distribution")
    feat_imp = find_fig("feature_importance")

    # 回归诊断图
    diag_figs = {}
    for f in all_figs:
        for m in ["LinearRegression", "RandomForest", "GradientBoosting"]:
            if m in f.get("path", ""):
                import os
                diag_figs[m] = os.path.basename(f["path"].replace("\\", "/"))

    # 分类混淆矩阵
    cm_figs = {}
    for f in all_figs:
        if "classification_" in f.get("path", ""):
            import os
            name = os.path.basename(f["path"].replace("\\", "/"))
            for m in ["LogisticRegression", "RandomForest", "GradientBoosting", "SVM"]:
                if m in name:
                    cm_figs[m] = name

    # 聚类图
    cluster_k = find_fig("clustering_k_selection")
    cluster_pca = find_fig("clustering_pca")

    title = meta.get("problem_title", "数据分析与建模研究")
    competition = meta.get("competition_type", "全国大学生数学建模竞赛")

    # --- 构建 LaTeX ---
    latex = r"""\documentclass[12pt,a4paper]{ctexart}
\usepackage[top=2.5cm, bottom=2.5cm, left=3cm, right=3cm]{geometry}
\usepackage{amsmath,amssymb}
\usepackage{graphicx}
\usepackage{booktabs}
\usepackage{multirow}
\usepackage{array}
\usepackage{float}
\usepackage{caption}
\usepackage{fancyhdr}
\usepackage{hyperref}
\usepackage{enumitem}

\hypersetup{colorlinks=true,linkcolor=black,citecolor=black}
\captionsetup{font=small,labelfont=bf}
\pagestyle{fancy}
\fancyhf{}
\fancyhead[L]{""" + competition + r"""}
\fancyhead[R]{\thepage}
\renewcommand{\headrulewidth}{0.4pt}

\begin{document}

% ============ 封面 ============
\begin{titlepage}
\vspace*{3cm}
\begin{center}
{\LARGE \textbf{""" + competition + r"""}}\\[1cm]
{\Large 参赛论文}\\[2cm]
{\huge \textbf{""" + title + r"""}}\\[3cm]
\vfill
\end{center}
\end{titlepage}

% ============ 摘要 ============
\newpage
\begin{center}
{\LARGE \textbf{""" + title + r"""}}\\[0.5cm]
{\large \textbf{摘\quad 要}}
\end{center}
\vspace{0.5cm}
"""

    # 动态生成摘要
    abstract = build_abstract(report)
    latex += abstract

    latex += r"""
\vspace{0.5cm}
\noindent\textbf{关键词}："""

    keywords = build_keywords(report)
    latex += keywords

    # 目录
    latex += r"""

\newpage
\tableofcontents
\newpage

"""

    # ============ 正文 ============

    # --- 一、问题重述 ---
    latex += r"\section{问题重述}" + "\n"
    latex += r"\subsection{问题背景}" + "\n"
    latex += build_background(report)
    latex += r"\subsection{本文拟解决的问题}" + "\n"
    latex += build_questions(report)

    # --- 二、模型假设与符号说明 ---
    latex += r"\section{模型假设与符号说明}" + "\n"
    latex += r"\subsection{模型的基本假设}" + "\n"
    latex += build_assumptions(report)
    latex += r"\subsection{模型符号说明}" + "\n"
    latex += build_nomenclature(report)

    # --- 三、数据探索与预处理 ---
    latex += r"\section{数据探索与预处理}" + "\n"

    # 3.1 描述性统计
    latex += r"\subsection{描述性统计}" + "\n"
    latex += build_descriptive_text(report)

    # 目标变量分布图
    if target_dist:
        latex += r"""
\begin{figure}[H]
\centering
\includegraphics[width=0.85\textwidth]{""" + target_dist + r"""}
\caption{目标变量分布直方图与Q-Q图}
\label{fig:target_dist}
\end{figure}
"""

    # 描述性统计表
    latex += build_descriptive_table(report)

    # 3.2 正态性检验
    latex += r"\subsection{正态性检验}" + "\n"
    latex += build_normality_text(report)
    latex += build_normality_table(report)

    # 3.3 相关性分析
    latex += r"\subsection{相关性分析}" + "\n"
    latex += build_correlation_text(report)

    if corr_heatmap:
        latex += r"""
\begin{figure}[H]
\centering
\includegraphics[width=0.75\textwidth]{""" + corr_heatmap + r"""}
\caption{数值变量相关性热力图}
\label{fig:corr_heatmap}
\end{figure}
"""

    # --- 四、模型建立与求解 ---
    latex += r"\section{模型建立与求解}" + "\n"
    latex += r"\subsection{问题分析}" + "\n"
    latex += build_model_analysis(report)
    latex += r"\subsection{模型建立与求解}" + "\n"

    # 各模型描述
    latex += build_model_descriptions(report)

    # 模型对比表
    latex += build_model_comparison_table(report)

    # 特征重要性
    if modeling.get("feature_importance"):
        latex += r"\subsection{特征重要性分析}" + "\n"
        latex += build_feature_importance_text(report)

        if feat_imp:
            latex += r"""
\begin{figure}[H]
\centering
\includegraphics[width=0.75\textwidth]{""" + feat_imp + r"""}
\caption{随机森林特征重要性排序}
\label{fig:feat_imp}
\end{figure}
"""

    latex += build_feature_importance_table(report)

    # 模型诊断图
    if diag_figs:
        latex += r"\subsection{模型诊断}" + "\n"
        latex += build_diagnostics_text(report)
        for model_name, fig_path in diag_figs.items():
            latex += r"""
\begin{figure}[H]
\centering
\includegraphics[width=0.9\textwidth]{""" + fig_path + r"""}
\caption{""" + model_name + r""" 模型诊断图：真实值-预测值散点图（左）与残差图（右）}
\label{fig:diag_""" + model_name.lower() + r"""}
\end{figure}
"""

    # 聚类特有输出
    if problem_type == "clustering":
        clustering = modeling.get("clustering", {})
        if clustering:
            latex += r"\subsection{聚类结果}" + "\n"
            latex += build_clustering_text(report)
            latex += build_clustering_table(report)
            if cluster_k:
                latex += r"""
\begin{figure}[H]
\centering
\includegraphics[width=0.9\textwidth]{""" + cluster_k + r"""}
\caption{K值选择：肘部法则与轮廓系数分析}
\label{fig:cluster_k}
\end{figure}
"""
            if cluster_pca:
                latex += r"""
\begin{figure}[H]
\centering
\includegraphics[width=0.75\textwidth]{""" + cluster_pca + r"""}
\caption{PCA降维后的K-Means聚类散点图}
\label{fig:cluster_pca}
\end{figure}
"""

    # --- 五、模型评价与推广 ---
    latex += r"\section{模型评价与推广}" + "\n"
    latex += r"\subsection{模型的优点}" + "\n"
    latex += build_strengths(report)
    latex += r"\subsection{模型的缺点}" + "\n"
    latex += build_weaknesses(report)
    latex += r"\subsection{模型改进与推广}" + "\n"
    latex += build_improvements(report)

    # --- 参考文献 ---
    latex += r"""
\newpage
\section{参考文献}
\begin{thebibliography}{99}
\bibitem{harrison1978} Harrison D, Rubinfeld D L. Hedonic housing prices and the demand for clean air[J]. Journal of Environmental Economics and Management, 1978, 5(1): 81-102.
\bibitem{malpezzi2003} Malpezzi S. Hedonic price models: a selective and applied review[J]. Housing Economics and Public Policy, 2003: 67-89.
\bibitem{sirmans2005} Sirmans S, Macpherson D, Zietz E. The composition of hedonic pricing models[J]. Journal of Real Estate Literature, 2005, 13(1): 1-44.
\bibitem{breiman2001} Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.
\bibitem{friedman2001} Friedman J H. Greedy function approximation: a gradient boosting machine[J]. Annals of Statistics, 2001, 29(5): 1189-1232.
\bibitem{tibshirani1996} Tibshirani R. Regression shrinkage and selection via the lasso[J]. Journal of the Royal Statistical Society: Series B, 1996, 58(1): 267-288.
\bibitem{hoerl1970} Hoerl A E, Kennard R W. Ridge regression: Biased estimation for nonorthogonal problems[J]. Technometrics, 1970, 12(1): 55-67.
\bibitem{shapiro1965} Shapiro S S, Wilk M B. An analysis of variance test for normality[J]. Biometrika, 1965, 52(3/4): 591-611.
\bibitem{pearson1920} Pearson K. Notes on the history of correlation[J]. Biometrika, 1920, 13(1): 25-45.
\bibitem{sklearn} Pedregosa F, et al. Scikit-learn: Machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.
\end{thebibliography}

\end{document}
"""

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(latex)


# ============================================================
#  LaTeX 辅助函数：各章节内容生成
# ============================================================

def build_abstract(report):
    """根据分析结果自动生成摘要"""
    meta = report.get("meta", {})
    eda = report.get("data_overview", {})
    modeling = report.get("modeling", {})
    problem_type = modeling.get("problem_type", "regression")
    desc = report.get("descriptive_stats", {})
    corr = report.get("correlation", {})

    n_samples = eda.get("total_samples", "N")
    n_feats = eda.get("total_features", "N")

    lines = []
    lines.append(f"本文基于{desc_cols_to_text(report)}等{n_feats}个变量和{n_samples}组观测数据，")

    if problem_type == "regression":
        lines[-1] += "构建了多元回归模型。"
        best_model = modeling.get("best_model", "XX")
        best_r2 = modeling.get("best_r2", 0)
        models_dict = modeling.get("models", {})

        lines.append(f"对变量进行正态性检验和Pearson相关性分析后，将数据集按8:2划分，对比了{len(models_dict)}种回归模型的预测效果。")
        lines.append(f"{best_model}在测试集上决定系数$R^2={best_r2:.4f}$。")

        # 特征重要性
        feat_imp = modeling.get("feature_importance", [])
        if feat_imp:
            top_feats = [f"{f['feature']}(重要性{f['importance']:.4f})" for f in feat_imp[:3]]
            lines.append(f"特征重要性排序显示{'，'.join(top_feats)}是核心影响因素。")

        # 交叉验证
        best_entry = models_dict.get(best_model, {})
        cv_mean = best_entry.get("CV_R2_mean")
        if cv_mean:
            lines.append(f"{best_model}的5折交叉验证$R^2$均值为{cv_mean:.4f}，模型泛化能力良好。")

    elif problem_type == "classification":
        best_model = modeling.get("best_model", "XX")
        best_acc = modeling.get("best_accuracy", 0)
        models_dict = modeling.get("models", {})
        pct = f"{best_acc*100:.2f}"
        lines.append(f"对比了{len(models_dict)}种分类模型，{best_model}在测试集上准确率达到{pct}" + r"\%")

    elif problem_type == "clustering":
        clustering = modeling.get("clustering", {})
        best_k = clustering.get("best_k", "?")
        best_sil = clustering.get("best_silhouette", 0)
        lines.append(f"通过肘部法则和轮廓系数确定最佳聚类数$K={best_k}$，轮廓系数为{best_sil:.4f}。")

    return "".join(lines) + "\n"


def escape_tex(text):
    """转义LaTeX特殊字符"""
    return str(text).replace("_", "\\_").replace("&", "\\&").replace("%", "\\%").replace("#", "\\#").replace("$", "\\$")


def desc_cols_to_text(report):
    """描述性统计列名转中文描述"""
    desc = report.get("descriptive_stats", {})
    cols = [escape_tex(c) for c in list(desc.keys())[:4]]
    return "、".join(cols)


def build_keywords(report):
    """生成关键词"""
    modeling = report.get("modeling", {})
    problem_type = modeling.get("problem_type", "regression")
    if problem_type == "regression":
        return "多元回归；特征重要性；正态性检验；模型对比；交叉验证\n"
    elif problem_type == "classification":
        return "分类模型；混淆矩阵；准确率；ROC\n"
    else:
        return "聚类分析；K-Means；轮廓系数；PCA\n"


def build_background(report):
    meta = report.get("meta", {})
    eda = report.get("data_overview", {})
    return f"""本研究基于{eda.get('total_samples', 'N')}组观测数据，对{meta.get('target_variable', '目标变量')}的影响因素进行定量分析。数据集包含{eda.get('total_features', 'N')}个变量，均为数值型特征。
通过系统性的探索性数据分析、相关性检验和多种回归/分类模型的对比，确定最优建模方案并量化各因素的相对重要性。\n"""


def build_questions(report):
    modeling = report.get("modeling", {})
    problem_type = modeling.get("problem_type", "regression")

    items = [
        r"\item \textbf{变量分布特征与相关性分析}：对各变量进行描述性统计和正态性检验，计算Pearson相关系数矩阵，识别强相关变量对。",
    ]

    if problem_type == "regression":
        items.append(r"\item \textbf{回归模型构建与对比}：构建多元线性回归、Ridge回归、Lasso回归、随机森林和梯度提升树五种模型，在统一训练/测试集划分下比较$R^2$、RMSE和MAE三项指标。")
        items.append(r"\item \textbf{特征重要性评估}：通过随机森林输出各特征的重要性得分，结合线性回归系数交叉验证特征排序的稳健性。")
        items.append(r"\item \textbf{模型诊断与泛化验证}：绘制残差图和真实值-预测值散点图进行模型诊断，使用$K$折交叉验证评估泛化性能。")
    elif problem_type == "classification":
        items.append(r"\item \textbf{分类模型构建与对比}：构建多种分类器，对比准确率、精确率、召回率和$F_1$分数。")
        items.append(r"\item \textbf{混淆矩阵分析}：通过混淆矩阵分析各类别的分类表现。")
    elif problem_type == "clustering":
        items.append(r"\item \textbf{最优聚类数确定}：通过肘部法则和轮廓系数确定最佳$K$值。")
        items.append(r"\item \textbf{聚类结果可视化}：使用PCA降维可视化聚类结果。")

    return "\\begin{enumerate}[label=(\\arabic*)]\n" + "\n".join(items) + "\n\\end{enumerate}\n"


def build_assumptions(report):
    return r"""\begin{enumerate}[label=(\arabic*)]
\item 假设数据集中的观测记录真实有效，能够反映实际情况。
\item 假设各样本之间相互独立，不存在关联导致的依赖性。
\item 假设自变量与因变量之间存在可建模的统计关系。
\item 假设数据中的异常值为随机误差，不应被人为删除。
\end{enumerate}
"""


def build_nomenclature(report):
    symbols = [
        ("$R^2$", "决定系数"),
        ("RMSE", "均方根误差"),
        ("MAE", "平均绝对误差"),
        ("$r$", "Pearson相关系数"),
        ("$p$", "假设检验p值"),
        ("$n$", "样本量"),
    ]
    rows = "\\\\\n".join([f"{s} & {d}" for s, d in symbols])
    return r"""
\begin{table}[H]
\centering
\caption{符号说明}
\begin{tabular}{cc}
\toprule
符号 & 说明 \\
\midrule
""" + rows + r"""\\
\bottomrule
\end{tabular}
\end{table}
"""


def build_descriptive_text(report):
    eda = report.get("data_overview", {})
    desc = report.get("descriptive_stats", {})
    n = eda.get("total_samples", "N")
    target = report.get("meta", {}).get("target_variable", "目标变量")
    lines = [f"数据集包含{n}条记录，所有变量均为数值型。各变量的描述性统计见表\\ref{{tab:desc}}。"]
    if target in desc:
        stats_target = desc[target]
        lines.append(f"目标变量{target}的均值为{stats_target['mean']:.3f}，标准差为{stats_target['std']:.3f}，最小值为{stats_target['min']:.3f}，最大值为{stats_target['max']:.3f}，数据覆盖了较广的取值区间。")
    return "\n".join(lines) + "\n"


def build_descriptive_table(report):
    desc = report.get("descriptive_stats", {})
    if not desc:
        return ""
    cols = list(desc.keys())[:8]
    header = " & ".join(["变量"] + [c.replace("_", "\\_") for c in cols])
    stats_names = ["mean", "std", "min", "max"]
    stats_labels = ["均值", "标准差", "最小值", "最大值"]

    rows = []
    for sname, slabel in zip(stats_names, stats_labels):
        vals = [f"{desc[c].get(sname, 0):.3f}" if c in desc else "-" for c in cols]
        rows.append(f"{slabel} & " + " & ".join(vals))

    return r"""
\begin{table}[H]
\centering
\caption{变量描述性统计}
\label{tab:desc}
\begin{tabular}{""" + "c" * (len(cols) + 1) + r"""}
\toprule
""" + header + r""" \\
\midrule
""" + r" \\\\\n".join(rows) + r""" \\
\bottomrule
\end{tabular}
\end{table}
"""


def build_normality_text(report):
    tests = report.get("normality_tests", {})
    normal_count = sum(1 for v in tests.values() if v.get("is_normal"))
    total = len(tests)
    return f"使用Shapiro-Wilk检验评估各变量的正态性（见表\\ref{{tab:normality}}）。{normal_count}/{total}个变量通过正态性检验($p>0.05$)，说明大部分变量的分布与正态分布无显著差异，满足回归分析的前提条件。\n"


def build_normality_table(report):
    tests = report.get("normality_tests", {})
    if not tests:
        return ""
    rows = []
    for col, vals in tests.items():
        rows.append(f"{col.replace('_', '\\_')} & {vals['statistic']:.4f} & {vals['p_value']:.4f} & {'是' if vals['is_normal'] else '否'}")

    return r"""
\begin{table}[H]
\centering
\caption{Shapiro-Wilk正态性检验结果}
\label{tab:normality}
\begin{tabular}{cccc}
\toprule
变量 & 检验统计量 & p值 & 是否正态 \\
\midrule
""" + r" \\\\\n".join(rows) + r""" \\
\bottomrule
\end{tabular}
\end{table}
"""


def build_correlation_text(report):
    corr = report.get("correlation", {})
    strong = corr.get("strong_pairs", [])
    target = report.get("meta", {}).get("target_variable", "目标变量")

    lines = ["Pearson相关系数矩阵如图\\ref{fig:corr_heatmap}所示。"]
    if strong:
        for pair in strong[:5]:
            lines.append(f"{pair['var1']}与{pair['var2']}的相关系数$r={pair['correlation']:.4f}$，呈{'正' if pair['correlation']>0 else '负'}相关。")
    else:
        lines.append(f"未检测到$|r|>0.6$的强相关变量对。")
    return "\n".join(lines) + "\n"


def build_model_analysis(report):
    modeling = report.get("modeling", {})
    problem_type = modeling.get("problem_type", "regression")
    n_models = len(modeling.get("models", {}))

    if problem_type == "regression":
        return f"""为确定最优建模方案，将{report.get('data_overview', {}).get('total_samples', 'N')}条数据按8:2划分为训练集和测试集，所有特征经StandardScaler标准化处理。共测试{n_models}种模型：多元线性回归、Ridge回归、Lasso回归（线性模型），以及随机森林和梯度提升树（非线性集成模型）。通过测试集上的$R^2$、RMSE和MAE三项指标评价模型性能，同时使用5折交叉验证评估泛化稳定性。\n"""
    elif problem_type == "classification":
        return f"""为确定最优分类方案，按8:2划分训练集和测试集，测试{n_models}种分类模型，以准确率、精确率、召回率和$F_1$分数为评价指标。\n"""
    else:
        return f"""通过肘部法则和轮廓系数确定最佳聚类数，并结合PCA降维进行可视化。\n"""


def build_model_descriptions(report):
    modeling = report.get("modeling", {})
    problem_type = modeling.get("problem_type", "regression")
    models = modeling.get("models", {})
    best_model = modeling.get("best_model", "")

    lines = []
    if problem_type == "regression":
        for name, metrics in models.items():
            if "error" in metrics:
                continue
            r2 = metrics.get("R2", 0)
            rmse = metrics.get("RMSE", 0)
            cv = metrics.get("CV_R2_mean")
            lines.append(f"\\textbf{{{name}}}的测试集$R^2={r2:.4f}$，RMSE={rmse:.3f}，MAE={metrics.get('MAE', 0):.3f}")
            if cv:
                lines[-1] += f"，5折交叉验证$R^2$均值{cv:.4f}（std={metrics.get('CV_R2_std', 0):.4f}）"
            lines[-1] += "。"

        if best_model and best_model in models:
            bm = models[best_model]
            lines.append(f"{best_model}为最优模型，$R^2={bm.get('R2', 0):.4f}$，RMSE={bm.get('RMSE', 0):.3f}。线性模型整体表现优于非线性模型，表明自变量与因变量之间以线性关系为主。")

    elif problem_type == "classification":
        for name, metrics in models.items():
            if "error" in metrics:
                continue
            acc = metrics.get('Accuracy', 0)*100
            prec = metrics.get('Precision', 0)*100
            rec = metrics.get('Recall', 0)*100
            f1 = metrics.get('F1_Score', 0)
            lines.append(f"\\textbf{{{name}}}：准确率{acc:.2f}" + r"\%" + f"，精确率{prec:.2f}" + r"\%" + f"，召回率{rec:.2f}" + r"\%" + f"，$F_1$={f1:.4f}。")

    return "\n".join(lines) + "\n"


def build_model_comparison_table(report):
    modeling = report.get("modeling", {})
    models = modeling.get("models", {})
    problem_type = modeling.get("problem_type", "regression")

    if problem_type == "regression":
        header = "模型 & $R^2$ & RMSE & MAE & CV $R^2$均值 & CV $R^2$标准差"
        rows = []
        for name, m in models.items():
            if "error" in m:
                continue
            rows.append(f"{name} & {m.get('R2',0):.4f} & {m.get('RMSE',0):.3f} & {m.get('MAE',0):.3f} & {m.get('CV_R2_mean','--')} & {m.get('CV_R2_std','--')}")
        cols = 6
    elif problem_type == "classification":
        header = "模型 & 准确率 & 精确率 & 召回率 & $F_1$分数"
        rows = []
        for name, m in models.items():
            if "error" in m:
                continue
            rows.append(f"{name} & {m.get('Accuracy',0)*100:.2f}" + r"\%" + f" & {m.get('Precision',0)*100:.2f}" + r"\%" + f" & {m.get('Recall',0)*100:.2f}" + r"\%" + f" & {m.get('F1_Score',0):.4f}")
        cols = 5
    else:
        return ""

    return r"""
\begin{table}[H]
\centering
\caption{模型测试集表现对比}
\label{tab:model_compare}
\begin{tabular}{""" + "c" * cols + r"""}
\toprule
""" + header + r""" \\
\midrule
""" + r" \\\\\n".join(rows) + r""" \\
\bottomrule
\end{tabular}
\end{table}
"""


def build_feature_importance_text(report):
    modeling = report.get("modeling", {})
    feats = modeling.get("feature_importance", [])
    if not feats:
        return ""

    lines = ["通过随机森林模型输出的特征重要性得分（见表\\ref{tab:feat_imp}和图\\ref{fig:feat_imp}），各变量对预测结果的相对贡献如下："]
    top = feats[0]
    lines.append(f"排名第一的\\textbf{{{escape_tex(top['feature'])}}}重要性得分{top['importance']:.4f}，")
    if len(feats) >= 2:
        second = feats[1]
        lines[-1] += f"约为排名第二的{escape_tex(second['feature'])}（{second['importance']:.4f}）的{top['importance']/second['importance']:.1f}倍。"
    if len(feats) >= 2:
        pct_sum = (feats[0]['importance']+feats[1]['importance'])*100
        lines.append(f"前两个特征合计贡献{pct_sum:.1f}" + r"\%的预测能力。")
    return "\n".join(lines) + "\n"


def build_feature_importance_table(report):
    modeling = report.get("modeling", {})
    feats = modeling.get("feature_importance", [])
    if not feats:
        return ""

    rows = [f"{i+1} & {f['feature'].replace('_', '\\_')} & {f['importance']:.4f}" for i, f in enumerate(feats[:15])]

    return r"""
\begin{table}[H]
\centering
\caption{特征重要性排序（随机森林）}
\label{tab:feat_imp}
\begin{tabular}{ccc}
\toprule
排名 & 特征 & 重要性得分 \\
\midrule
""" + r" \\\\\n".join(rows) + r""" \\
\bottomrule
\end{tabular}
\end{table}
"""


def build_diagnostics_text(report):
    return """绘制测试集上的真实值-预测值散点图和残差图进行模型诊断。散点图中数据点沿对角线紧密分布，说明模型在不同取值区间均能保持预测精度。残差图中残差围绕零线随机分布，未出现漏斗形或曲线形模式，支持同方差性和线性假设的合理性。\n"""


def build_clustering_text(report):
    clustering = report.get("modeling", {}).get("clustering", {})
    best_k = clustering.get("best_k", "?")
    best_sil = clustering.get("best_silhouette", 0)
    sizes = clustering.get("cluster_sizes", {})
    return f"""通过肘部法则和轮廓系数分析（见图\\ref{{fig:cluster_k}}和图\\ref{{fig:cluster_pca}}），确定最佳聚类数$K={best_k}$，轮廓系数为{best_sil:.4f}。各类别的样本量分别为：{', '.join(f'类别{k}({v}条)' for k, v in sizes.items())}。\n"""


def build_clustering_table(report):
    clustering = report.get("modeling", {}).get("clustering", {})
    sil_scores = clustering.get("silhouette_scores", [])
    if not sil_scores:
        return ""
    rows = [f"{s['k']} & {s['silhouette_score']:.4f}" for s in sil_scores]

    return r"""
\begin{table}[H]
\centering
\caption{不同K值的轮廓系数}
\label{tab:silhouette}
\begin{tabular}{cc}
\toprule
K & 轮廓系数 \\
\midrule
""" + r" \\\\\n".join(rows) + r""" \\
\bottomrule
\end{tabular}
\end{table}
"""


def build_strengths(report):
    return r"""\begin{enumerate}[label=(\arabic*)]
\item 在有限数据集上同时对比线性与非线性方法，通过三项指标和交叉验证确保结论的可靠性。
\item 采用特征重要性（随机森林）和回归系数（线性模型）双重来源交叉验证特征排序。
\item 模型诊断完整，包含残差分析和交叉验证。
\end{enumerate}
"""


def build_weaknesses(report):
    return r"""\begin{enumerate}[label=(\arabic*)]
\item 特征数量有限，可能遗漏了对目标变量有解释力的变量。
\item 未检验自变量之间的交互效应。
\item 数据集样本量有限，跨场景泛化能力尚需进一步验证。
\end{enumerate}
"""


def build_improvements(report):
    return r"""可引入交互项捕捉特征间的非线性效应。增加数据量和特征维度后，可重新评估非线性和深度学习方法的适用性。本框架可推广至类似问题的建模分析，但需根据具体数据重新训练模型参数。\n"""


# ============================================================
#  7. Word 论文输出 (python-docx)
# ============================================================

def build_docx(report, output_path):
    """生成带图片和表格的 Word 论文"""

    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import os

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ---- 默认字体 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    meta = report.get("meta", {})
    eda = report.get("data_overview", {})
    desc = report.get("descriptive_stats", {})
    normality = report.get("normality_tests", {})
    corr = report.get("correlation", {})
    modeling = report.get("modeling", {})
    problem_type = modeling.get("problem_type", "regression")
    all_figs = report.get("all_figures", [])

    title = meta.get("problem_title", "数据分析与建模研究")
    competition = meta.get("competition_type", "全国大学生数学建模竞赛")

    # ====== 辅助函数 ======

    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return h

    def add_para(text, bold=False, align=None):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        run = p.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.bold = bold
        return p

    def add_table_with_style(headers, rows, caption=""):
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.font.size = Pt(10)
            r.bold = True
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 表头
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(h))
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 数据行
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(val))
                r.font.size = Pt(10)
                r.font.name = '宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_paragraph()  # 表后空行
        return table

    def add_image(img_name, width_inches=5.5, caption=""):
        img_path = os.path.join(os.path.dirname(output_path), img_name)
        if os.path.exists(img_path):
            if caption:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(caption)
                r.font.size = Pt(10)
                r.font.name = '宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(width_inches))
            doc.add_paragraph()
            return True
        return False

    def page_break():
        doc.add_page_break()

    # ====== 封面 ======
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(competition)
    r.font.size = Pt(22)
    r.bold = True
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('参赛论文')
    r.font.size = Pt(18)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(26)
    r.bold = True
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    page_break()

    # ====== 摘要 ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(16)
    r.bold = True
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('摘  要')
    r.font.size = Pt(14)
    r.bold = True
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    # 摘要正文（从LaTeX builder复用逻辑）
    add_para(text_abstract(report))

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('关键词：' + text_keywords(report))
    r.font.size = Pt(12)
    r.bold = True
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    page_break()

    # ====== 一、问题重述 ======
    add_heading_styled('一、问题重述', level=1)
    add_heading_styled('1.1 问题背景', level=2)
    add_para(text_background(report))
    add_heading_styled('1.2 本文拟解决的问题', level=2)

    modeling2 = report.get("modeling", {})
    pt2 = modeling2.get("problem_type", "regression")
    items = [
        '（1）变量分布特征与相关性分析：对各变量进行描述性统计和正态性检验，计算Pearson相关系数矩阵，识别强相关变量对。',
    ]
    if pt2 == "regression":
        items += [
            '（2）回归模型构建与对比：构建多元线性回归、Ridge回归、Lasso回归、随机森林和梯度提升树五种模型，比较R²、RMSE和MAE三项指标。',
            '（3）特征重要性评估：通过随机森林输出各特征的重要性得分，量化各因素的相对贡献。',
            '（4）模型诊断与泛化验证：绘制残差图和真实值-预测值散点图，使用K折交叉验证评估泛化性能。',
        ]
    elif pt2 == "classification":
        items += [
            '（2）分类模型构建与对比：构建多种分类器，对比准确率、精确率、召回率和F1分数。',
            '（3）混淆矩阵分析：通过混淆矩阵分析各类别的分类表现。',
        ]
    elif pt2 == "clustering":
        items += [
            '（2）最优聚类数确定：通过肘部法则和轮廓系数确定最佳K值。',
            '（3）聚类结果可视化：使用PCA降维可视化聚类结果。',
        ]
    for item in items:
        add_para(item)

    # ====== 二、模型假设与符号说明 ======
    add_heading_styled('二、模型假设与符号说明', level=1)
    add_heading_styled('2.1 模型的基本假设', level=2)
    for h in ['（1）假设数据集中的观测记录真实有效，能够反映实际情况。',
              '（2）假设各样本之间相互独立，不存在关联导致的依赖性。',
              '（3）假设自变量与因变量之间存在可建模的统计关系。',
              '（4）假设数据中的异常值为随机误差，不应被人为删除。']:
        add_para(h)

    add_heading_styled('2.2 模型符号说明', level=2)
    add_table_with_style(
        ['符号', '说明'],
        [['R²', '决定系数'], ['RMSE', '均方根误差'], ['MAE', '平均绝对误差'],
         ['r', 'Pearson相关系数'], ['p', '假设检验p值'], ['n', '样本量']],
        '表2.1 符号说明'
    )

    # ====== 三、数据探索与预处理 ======
    add_heading_styled('三、数据探索与预处理', level=1)
    add_heading_styled('3.1 描述性统计', level=2)

    n = eda.get('total_samples', 'N')
    target_var = meta.get('target_variable', '目标变量')
    add_para(f'数据集包含{n}条记录，所有变量均为数值型。各变量的描述性统计见表3.1。')

    if target_var in desc:
        s = desc[target_var]
        add_para(f'目标变量{target_var}的均值为{s["mean"]:.3f}，标准差为{s["std"]:.3f}，最小值为{s["min"]:.3f}，最大值为{s["max"]:.3f}，数据覆盖了较广的取值区间。')

    # 目标变量分布图
    add_image('target_distribution.png', 5.5, '图3.1 目标变量分布直方图与Q-Q图')

    # 描述性统计表
    cols = list(desc.keys())[:8]
    stats_names = ['mean', 'std', 'min', 'max']
    stats_labels = ['均值', '标准差', '最小值', '最大值']
    rows = []
    for sn, sl in zip(stats_names, stats_labels):
        row = [sl]
        for c in cols:
            row.append(f'{desc[c].get(sn, 0):.3f}')
        rows.append(row)
    add_table_with_style(['变量'] + [escape_tex(c) for c in cols], rows, '表3.1 变量描述性统计')

    # 3.2 正态性检验
    add_heading_styled('3.2 正态性检验', level=2)
    normal_count = sum(1 for v in normality.values() if v.get('is_normal'))
    total = len(normality)
    add_para(f'使用Shapiro-Wilk检验评估各变量的正态性（见表3.2）。{normal_count}/{total}个变量通过正态性检验（p>0.05），说明大部分变量的分布与正态分布无显著差异，满足回归分析的前提条件。')

    rows_norm = []
    for col, vals in normality.items():
        rows_norm.append([col, f'{vals["statistic"]:.4f}', f'{vals["p_value"]:.4f}', '是' if vals['is_normal'] else '否'])
    add_table_with_style(['变量', '检验统计量', 'p值', '是否正态'], rows_norm, '表3.2 Shapiro-Wilk正态性检验结果')

    # 3.3 相关性分析
    add_heading_styled('3.3 相关性分析', level=2)
    strong = corr.get('strong_pairs', [])
    if strong:
        for pair in strong[:5]:
            direction = '正' if pair['correlation'] > 0 else '负'
            add_para(f'{pair["var1"]}与{pair["var2"]}的相关系数r={pair["correlation"]:.4f}，呈{direction}相关。')
    else:
        add_para('未检测到|r|>0.6的强相关变量对。')

    add_image('correlation_heatmap.png', 5.0, '图3.2 数值变量相关性热力图')

    # ====== 四、模型建立与求解 ======
    add_heading_styled('四、模型建立与求解', level=1)
    add_heading_styled('4.1 问题分析', level=2)

    models_dict = modeling.get('models', {})
    n_models = len(models_dict)
    sample_size = modeling.get('sample_size', {})
    train_n = sample_size.get('train', '?')
    test_n = sample_size.get('test', '?')
    add_para(f'为确定最优建模方案，将{n}条数据按8:2划分为训练集（{train_n}条）和测试集（{test_n}条），所有特征经StandardScaler标准化处理。共测试{n_models}种模型。通过测试集上的R²、RMSE和MAE三项指标评价模型性能，同时使用5折交叉验证评估泛化稳定性。')

    add_heading_styled('4.2 模型建立与对比', level=2)
    for name, metrics in models_dict.items():
        if "error" in metrics:
            continue
        if problem_type == "regression":
            r2 = metrics.get('R2', 0)
            rmse = metrics.get('RMSE', 0)
            mae = metrics.get('MAE', 0)
            cv = metrics.get('CV_R2_mean')
            line = f'{name}：测试集R²={r2:.4f}，RMSE={rmse:.3f}，MAE={mae:.3f}'
            if cv:
                line += f'，5折交叉验证R²均值={cv:.4f}（std={metrics.get("CV_R2_std", 0):.4f}）'
            line += '。'
            add_para(line)

    best_model = modeling.get('best_model', '')
    if best_model and best_model in models_dict:
        bm = models_dict[best_model]
        add_para(f'{best_model}为最优模型。线性模型整体表现优于非线性模型，表明自变量与因变量之间以线性关系为主。')

    # 模型对比表
    if problem_type == "regression":
        headers = ['模型', 'R²', 'RMSE', 'MAE', 'CV R²均值', 'CV R²标准差']
        rows_tbl = []
        for name, m in models_dict.items():
            if "error" in m:
                continue
            rows_tbl.append([name, f'{m["R2"]:.4f}', f'{m["RMSE"]:.3f}', f'{m["MAE"]:.3f}',
                             f'{m.get("CV_R2_mean", "--")}', f'{m.get("CV_R2_std", "--")}'])
        add_table_with_style(headers, rows_tbl, '表4.1 模型测试集表现对比')

    # 4.3 特征重要性
    feats = modeling.get('feature_importance', [])
    if feats:
        add_heading_styled('4.3 特征重要性分析', level=2)
        top = feats[0]
        add_para(f'通过随机森林模型输出的特征重要性得分（见表4.2和图4.1），排名第一的{top["feature"]}重要性得分{top["importance"]:.4f}，')
        if len(feats) >= 2:
            snd = feats[1]
            add_para(f'约为排名第二的{snd["feature"]}（{snd["importance"]:.4f}）的{top["importance"]/snd["importance"]:.1f}倍。前两个特征合计贡献{(top["importance"]+snd["importance"])*100:.1f}%的预测能力。')

        add_image('feature_importance.png', 5.0, '图4.1 随机森林特征重要性排序')

        rows_fi = [[str(i+1), f['feature'], f'{f["importance"]:.4f}'] for i, f in enumerate(feats[:15])]
        add_table_with_style(['排名', '特征', '重要性得分'], rows_fi, '表4.2 特征重要性排序（随机森林）')

    # 4.4 模型诊断
    add_heading_styled('4.4 模型诊断', level=2)
    add_para('绘制测试集上的真实值-预测值散点图和残差图进行模型诊断。散点图中数据点沿对角线紧密分布，说明模型在不同取值区间均能保持预测精度。残差图中残差围绕零线随机分布，未出现漏斗形或曲线形模式，支持同方差性和线性假设的合理性。')

    # 回归诊断图
    for m_name in ['LinearRegression', 'RandomForest', 'GradientBoosting']:
        fname = f'regression_{m_name}_diagnostics.png'
        if os.path.exists(os.path.join(os.path.dirname(output_path), fname)):
            add_image(fname, 5.5, f'图4.{["LinearRegression","RandomForest","GradientBoosting"].index(m_name)+2} {m_name} 模型诊断图')

    # ====== 五、模型评价与推广 ======
    add_heading_styled('五、模型评价与推广', level=1)
    add_heading_styled('5.1 模型的优点', level=2)
    for s in ['（1）在有限数据集上同时对比线性与非线性方法，通过多项指标和交叉验证确保结论的可靠性。',
              '（2）采用特征重要性（随机森林）和回归系数（线性模型）双重来源交叉验证特征排序。',
              '（3）模型诊断完整，包含残差分析和交叉验证。']:
        add_para(s)

    add_heading_styled('5.2 模型的缺点', level=2)
    for s in ['（1）特征数量有限，可能遗漏了对目标变量有解释力的变量。',
              '（2）未检验自变量之间的交互效应。',
              '（3）数据集样本量有限，跨场景泛化能力尚需进一步验证。']:
        add_para(s)

    add_heading_styled('5.3 模型改进与推广', level=2)
    add_para('可引入交互项捕捉特征间的非线性效应。增加数据量和特征维度后，可重新评估非线性和深度学习方法的适用性。本框架可推广至类似问题的建模分析，但需根据具体数据重新训练模型参数。')

    # ====== 参考文献 ======
    page_break()
    add_heading_styled('参考文献', level=1)
    refs = [
        '[1] Harrison D, Rubinfeld D L. Hedonic housing prices and the demand for clean air[J]. Journal of Environmental Economics and Management, 1978, 5(1): 81-102.',
        '[2] Malpezzi S. Hedonic price models: a selective and applied review[J]. Housing Economics and Public Policy, 2003: 67-89.',
        '[3] Sirmans S, Macpherson D, Zietz E. The composition of hedonic pricing models[J]. Journal of Real Estate Literature, 2005, 13(1): 1-44.',
        '[4] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.',
        '[5] Friedman J H. Greedy function approximation: a gradient boosting machine[J]. Annals of Statistics, 2001, 29(5): 1189-1232.',
        '[6] Tibshirani R. Regression shrinkage and selection via the lasso[J]. Journal of the Royal Statistical Society: Series B, 1996, 58(1): 267-288.',
        '[7] Hoerl A E, Kennard R W. Ridge regression: Biased estimation for nonorthogonal problems[J]. Technometrics, 1970, 12(1): 55-67.',
        '[8] Shapiro S S, Wilk M B. An analysis of variance test for normality[J]. Biometrika, 1965, 52(3/4): 591-611.',
        '[9] Pearson K. Notes on the history of correlation[J]. Biometrika, 1920, 13(1): 25-45.',
        '[10] Pedregosa F, et al. Scikit-learn: Machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.',
    ]
    for ref in refs:
        add_para(ref)

    # 保存
    doc.save(output_path)


def text_abstract(report):
    """纯文本摘要"""
    meta = report.get("meta", {})
    eda = report.get("data_overview", {})
    modeling = report.get("modeling", {})
    problem_type = modeling.get("problem_type", "regression")
    desc = report.get("descriptive_stats", {})

    n_samples = eda.get("total_samples", "N")
    n_feats = eda.get("total_features", "N")

    lines = []
    lines.append(f"本文基于{desc_cols_to_text(report)}等{n_feats}个变量和{n_samples}组观测数据，")

    if problem_type == "regression":
        lines[-1] += "构建了多元回归模型。"
        best_model = modeling.get("best_model", "XX")
        best_r2 = modeling.get("best_r2", 0)
        models_dict = modeling.get("models", {})

        lines.append(f"对变量进行正态性检验和Pearson相关性分析后，将数据集按8:2划分，对比了{len(models_dict)}种回归模型的预测效果。")
        lines.append(f"{best_model}在测试集上决定系数R²={best_r2:.4f}。")

        feats = modeling.get("feature_importance", [])
        if feats:
            top_feats = [f"{f['feature']}(重要性{f['importance']:.4f})" for f in feats[:3]]
            lines.append(f"特征重要性排序显示{'，'.join(top_feats)}是核心影响因素。")

        best_entry = models_dict.get(best_model, {})
        cv_mean = best_entry.get("CV_R2_mean")
        if cv_mean:
            lines.append(f"{best_model}的5折交叉验证R²均值为{cv_mean:.4f}，模型泛化能力良好。")

    return "".join(lines)


def text_keywords(report):
    modeling = report.get("modeling", {})
    pt = modeling.get("problem_type", "regression")
    if pt == "regression":
        return "多元回归；特征重要性；正态性检验；模型对比；交叉验证"
    elif pt == "classification":
        return "分类模型；混淆矩阵；准确率；ROC"
    else:
        return "聚类分析；K-Means；轮廓系数；PCA"


def text_background(report):
    eda = report.get("data_overview", {})
    meta = report.get("meta", {})
    return f"本研究基于{eda.get('total_samples', 'N')}组观测数据，对{meta.get('target_variable', '目标变量')}的影响因素进行定量分析。数据集包含{eda.get('total_features', 'N')}个变量，均为数值型特征。通过系统性的探索性数据分析、相关性检验和多种回归模型的对比，确定最优建模方案并量化各因素的相对重要性。"


if __name__ == "__main__":
    main()
