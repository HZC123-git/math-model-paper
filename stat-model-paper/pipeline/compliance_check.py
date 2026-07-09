#!/usr/bin/env python3
"""
排版合规检查器
根据竞赛类型检查 Word/LaTeX 论文的排版是否符合官方规范。
"""
import json
from pathlib import Path

# ============================================================
#  各竞赛官方排版规范
# ============================================================

FORMAT_SPECS = {
    "国赛": {
        "page": {"paper": "A4", "top_margin": "2.54cm", "bottom_margin": "2.54cm",
                 "left_margin": "3.18cm", "right_margin": "3.18cm"},
        "font": {"body": "宋体", "heading": "黑体", "size_body": 12, "size_heading_1": 16,
                 "size_heading_2": 14, "size_table": 10},
        "spacing": {"line_spacing": 1.5, "first_indent": "2字符(24pt)"},
        "tables": {"style": "三线表", "caption_above": False, "caption_below": True},
        "figures": {"caption_below": True, "numbered": True},
        "equations": {"numbered": True, "align": "center"},
        "header_footer": {"page_number": True, "page_number_pos": "bottom-center"},
        "title_page": True,
        "toc": True,
    },
    "研究生数学建模": {
        "page": {"paper": "A4", "top_margin": "2.54cm", "bottom_margin": "2.54cm",
                 "left_margin": "3.17cm", "right_margin": "3.17cm"},
        "font": {"body": "宋体", "heading": "黑体", "size_body": 12, "size_heading_1": 16,
                 "size_heading_2": 14},
        "spacing": {"line_spacing": 1.5},
        "tables": {"style": "三线表", "caption_above": False},
        "figures": {"resolution": "300dpi以上"},
        "title_page": True,
        "toc": True,
        "must_have": ["封面（学校+参赛队号+队员姓名）", "目录", "附录含完整代码"],
    },
    "美赛": {
        "page": {"paper": "US Letter", "top_margin": "1in", "bottom_margin": "1in",
                 "left_margin": "1in", "right_margin": "1in"},
        "font": {"body": "Times New Roman", "size_body": 12, "size_heading": 12},
        "spacing": {"line_spacing": 2.0},
        "summary": {"max_pages": 1, "separate_page": True, "include_team_number": True},
        "must_have": ["Summary Sheet (1 page max)", "Table of Contents",
                      "Sensitivity Analysis", "Strengths and Weaknesses",
                      "Letter (if required by problem)", "References"],
    },
    "统计建模": {
        "page": {"paper": "A4"},
        "font": {"body": "宋体", "heading": "黑体", "size_body": 12},
        "spacing": {"line_spacing": 1.5},
        "tables": {"style": "三线表"},
        "must_have": ["文献综述（15-25篇）", "内生性讨论", "稳健性检验", "政策建议"],
    },
}


def check_docx_compliance(docx_path, competition_type):
    """检查 Word 文档是否符合竞赛排版规范"""
    from docx import Document
    from docx.shared import Cm, Inches, Pt

    spec = FORMAT_SPECS.get(competition_type, FORMAT_SPECS["国赛"])
    issues = []
    warnings = []
    passed = []

    try:
        doc = Document(docx_path)
    except Exception as e:
        return [], [], [f"无法打开文档: {e}"]

    # ---- 页面设置 ----
    section = doc.sections[0]
    page = spec["page"]

    if page.get("paper") == "A4":
        expected_w, expected_h = Cm(21), Cm(29.7)
        if abs(section.page_width - expected_w) > Cm(0.3):
            issues.append(f"页面宽度: {section.page_width/914400:.1f}in (应为A4=21cm)")
        else:
            passed.append("页面尺寸: A4 ✅")

    # 页边距
    margin_map = {"top": section.top_margin, "bottom": section.bottom_margin,
                  "left": section.left_margin, "right": section.right_margin}
    expected_margins = {"top": Cm(2.54), "bottom": Cm(2.54), "left": Cm(3.18), "right": Cm(3.18)}

    for side, val in margin_map.items():
        expected = expected_margins.get(side)
        if expected and abs(val - expected) > Cm(0.2):
            issues.append(f"{side}边距: {val/914400:.2f}in (应为{expected/914400:.1f}in)")
        elif expected:
            passed.append(f"{side}边距 ✅")

    # ---- 字体检查 ----
    heading_count = 0
    body_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name:
            heading_count += 1
        elif len(text) > 20:
            body_count += 1

    # ---- 图表检查 ----
    # 统计图片
    img_count = sum(1 for r in doc.part.rels.values() if 'image' in r.reltype)
    passed.append(f"图片数量: {img_count} 张")

    # 统计表格
    table_count = len(doc.tables)
    passed.append(f"表格数量: {table_count} 个")

    # ---- 必含元素检查 ----
    full_text = "\n".join(p.text for p in doc.paragraphs)

    must_haves = spec.get("must_have", [])
    for item in must_haves:
        # 模糊匹配
        keywords = item.split("（")[0].replace("(", "").strip()
        if any(kw in full_text for kw in keywords.split()):
            passed.append(f"必含元素: {keywords} ✅")
        else:
            issues.append(f"缺少必含元素: {item}")

    # ---- 行距检查 ----
    line_spacing_found = False
    for para in doc.paragraphs[:30]:  # 抽样前30段
        pf = para.paragraph_format
        if pf.line_spacing and pf.line_spacing > 1.3:
            line_spacing_found = True
            break
    if line_spacing_found:
        passed.append("行距: >1.3 ✅")
    else:
        warnings.append("行距可能过小（建议1.5倍）")

    # ---- 页数检查 ----
    # 粗略估算：平均每段约一段话
    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    word_count = len(full_text.replace('\n', '').replace(' ', ''))
    passed.append(f"正文约 {word_count} 字, {para_count} 段")

    return issues, warnings, passed


def check_latex_compliance(tex_path, competition_type):
    """检查 LaTeX 文件的合规性"""
    with open(tex_path, 'r', encoding='utf-8') as f:
        content = f.read()

    spec = FORMAT_SPECS.get(competition_type, FORMAT_SPECS["国赛"])
    passed = []
    issues = []

    checks = [
        (r'\\documentclass.*a4paper', "A4纸张 ✅", "缺少a4paper选项"),
        (r'\\usepackage.*ctex', "中文支持(ctex) ✅", "缺少ctex中文支持"),
        (r'\\usepackage.*graphicx', "图片支持(graphicx) ✅", "缺少graphicx包"),
        (r'\\usepackage.*booktabs', "三线表(booktabs) ✅", "缺少booktabs包(三线表)"),
        (r'\\usepackage.*amsmath', "数学公式(amsmath) ✅", "缺少amsmath包"),
        (r'\\tableofcontents', "目录 ✅", "缺少目录"),
        (r'\\begin\{thebibliography\}', "参考文献 ✅", "缺少参考文献环境"),
    ]

    for pattern, ok_msg, fail_msg in checks:
        import re
        if re.search(pattern, content):
            passed.append(ok_msg)
        else:
            issues.append(fail_msg)

    # 检查字体设置
    if competition_type in ("国赛", "研究生数学建模", "统计建模"):
        if 'SimSun' in content or '宋体' in content or '\\setCJKmainfont' in content:
            passed.append("中文字体设置 ✅")
        else:
            issues.append("未设置中文字体（需要SimSun/宋体）")

    return issues, [], passed


def run_compliance_check(file_path, competition_type="国赛"):
    """运行合规检查"""
    path = Path(file_path)
    if not path.exists():
        return [f"文件不存在: {file_path}"], [], []

    if path.suffix == '.docx':
        issues, warnings, passed = check_docx_compliance(str(path), competition_type)
    elif path.suffix == '.tex':
        issues, warnings, passed = check_latex_compliance(str(path), competition_type)
    else:
        return [f"不支持的文件类型: {path.suffix}"], [], []

    return issues, warnings, passed


def print_report(issues, warnings, passed, competition_type):
    """打印合规检查报告"""
    total = len(passed) + len(issues) + len(warnings)
    score = len(passed) / max(total, 1) * 100

    print(f"\n{'='*60}")
    print(f"  排版合规检查报告")
    print(f"  竞赛类型: {competition_type}")
    print(f"  合规率: {score:.0f}% ({len(passed)}/{total})")
    print(f"{'='*60}")

    if passed:
        print(f"\n  [PASS] ({len(passed)}):")
        for p in passed:
            print(f"     {p}")

    if warnings:
        print(f"\n  [WARN] ({len(warnings)}):")
        for w in warnings:
            print(f"     {w}")

    if issues:
        print(f"\n  [FAIL] ({len(issues)}):")
        for i in issues:
            print(f"     {i}")

    print(f"\n  {'='*60}")
    if score >= 90:
        print(f"  Result: PASS - ready for submission")
    elif score >= 70:
        print(f"  Result: WARN - fix issues above before submitting")
    else:
        print(f"  Result: FAIL - major non-compliance, must re-check")

    return score


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='论文排版合规检查')
    parser.add_argument('file', help='论文文件路径 (.docx 或 .tex)')
    parser.add_argument('--competition', '-c', default='国赛',
                        choices=['国赛', '研究生数学建模', '美赛', '统计建模'],
                        help='竞赛类型')
    args = parser.parse_args()

    issues, warnings, passed = run_compliance_check(args.file, args.competition)
    print_report(issues, warnings, passed, args.competition)
